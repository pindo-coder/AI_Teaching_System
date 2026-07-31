from __future__ import annotations

import json
import os
from pathlib import Path
import re
import shutil
import subprocess
import tempfile
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.core.config import BACKEND_DIR, settings


PRESENTATION_RUNTIME = BACKEND_DIR / "presentation_runtime"
PRESENTATION_SCRIPT = PRESENTATION_RUNTIME / "render_pptx.mjs"


def _safe_name(value: str, fallback: str) -> str:
    cleaned = re.sub(r'[\\/:*?"<>|\r\n]+', "_", value).strip(" ._")
    return cleaned[:80] or fallback


class PresentationArtifactService:
    """将结构化教学成果渲染为可下载文件。"""

    def __init__(self, run_id: int) -> None:
        self.run_id = run_id
        root = Path(settings.generated_artifact_directory)
        if not root.is_absolute():
            root = (BACKEND_DIR / root).resolve()
        self.root = root
        self.run_dir = self.root / str(run_id)
        self.run_dir.mkdir(parents=True, exist_ok=True)

    def _node_binary(self) -> str:
        configured = settings.presentation_node_binary
        resolved = shutil.which(configured)
        if resolved:
            return resolved
        path = Path(configured)
        if path.exists():
            return str(path)
        raise RuntimeError("未找到 Node.js，请安装 Node 20+ 或配置 PRESENTATION_NODE_BINARY")

    def render_pptx(
        self,
        *,
        course_name: str,
        chapter_title: str,
        ppt_data: dict[str, Any],
        evidence: list[dict[str, Any]],
    ) -> dict[str, Any]:
        if not PRESENTATION_SCRIPT.exists():
            raise RuntimeError("PPT 生成脚本不存在")
        file_stem = _safe_name(ppt_data.get("title") or chapter_title, "教学课件")
        output_path = self.run_dir / f"{file_stem}.pptx"
        payload = {
            "course_name": course_name,
            "chapter_title": chapter_title,
            "ppt": ppt_data,
            "evidence": evidence,
            "artifact_root": str(self.root.resolve()),
        }
        with tempfile.NamedTemporaryFile(
            mode="w",
            encoding="utf-8",
            suffix=".json",
            dir=self.run_dir,
            delete=False,
        ) as handle:
            json.dump(payload, handle, ensure_ascii=False)
            input_path = Path(handle.name)
        env = os.environ.copy()
        module_paths: list[str] = []
        local_modules = PRESENTATION_RUNTIME / "node_modules"
        if local_modules.exists():
            module_paths.append(str(local_modules))
        if settings.presentation_node_modules:
            module_paths.append(settings.presentation_node_modules)
        if env.get("NODE_PATH"):
            module_paths.append(env["NODE_PATH"])
        if module_paths:
            env["NODE_PATH"] = os.pathsep.join(module_paths)
        try:
            result = subprocess.run(
                [self._node_binary(), str(PRESENTATION_SCRIPT), str(input_path), str(output_path)],
                cwd=PRESENTATION_RUNTIME,
                env=env,
                capture_output=True,
                text=True,
                timeout=120,
                check=False,
            )
        finally:
            input_path.unlink(missing_ok=True)
        if result.returncode != 0 or not output_path.exists():
            detail = (result.stderr or result.stdout or "PPT 渲染失败").strip()
            raise RuntimeError(f"PPT 渲染失败：{detail[-800:]}")
        return {
            "kind": "ppt",
            "title": ppt_data.get("title") or f"{chapter_title}教学课件",
            "file_name": output_path.name,
            "storage_path": str(output_path.relative_to(self.root)),
            "media_type": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
            "slide_count": len(ppt_data.get("slides") or []),
            "preview": ppt_data,
        }

    def render_lesson_plan(
        self,
        *,
        course_name: str,
        chapter_title: str,
        lesson_plan: dict[str, Any],
        evidence: list[dict[str, Any]],
    ) -> dict[str, Any]:
        title = lesson_plan.get("title") or f"{chapter_title}教学教案"
        output_path = self.run_dir / f"{_safe_name(title, '教学教案')}.docx"
        document = Document()
        document.styles["Normal"].font.name = "宋体"
        document.styles["Normal"].font.size = Pt(11)
        heading = document.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph(f"课程：{course_name}")
        document.add_paragraph(f"专题：{chapter_title}")
        document.add_heading("一、课程定位", level=1)
        document.add_paragraph(str(lesson_plan.get("overview") or ""))
        document.add_heading("二、教学目标", level=1)
        for item in lesson_plan.get("objectives") or []:
            document.add_paragraph(str(item), style="List Bullet")
        document.add_heading("三、课前准备", level=1)
        for item in lesson_plan.get("preparation") or []:
            document.add_paragraph(str(item), style="List Bullet")
        document.add_heading("四、教学过程", level=1)
        table = document.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        for cell, text in zip(
            table.rows[0].cells,
            ["教学环节", "时间", "教师活动", "学生活动", "资料依据"],
        ):
            cell.text = text
        for item in lesson_plan.get("procedures") or []:
            cells = table.add_row().cells
            cells[0].text = str(item.get("stage") or "")
            cells[1].text = f"{item.get('duration_minutes') or 0} 分钟"
            cells[2].text = str(item.get("teacher_activity") or "")
            cells[3].text = str(item.get("student_activity") or "")
            cells[4].text = "、".join(item.get("evidence_refs") or [])
        document.add_heading("五、学习评价", level=1)
        for item in lesson_plan.get("assessment") or []:
            document.add_paragraph(str(item), style="List Bullet")
        document.add_heading("六、课后任务", level=1)
        document.add_paragraph(str(lesson_plan.get("homework") or ""))
        self._append_sources(document, evidence)
        document.save(output_path)
        return {
            "kind": "lesson_plan",
            "title": title,
            "file_name": output_path.name,
            "storage_path": str(output_path.relative_to(self.root)),
            "media_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "preview": lesson_plan,
        }

    def render_activity_guide(
        self,
        *,
        course_name: str,
        chapter_title: str,
        activities: list[dict[str, Any]],
        evidence: list[dict[str, Any]],
    ) -> dict[str, Any]:
        title = f"{chapter_title}课堂活动设计"
        output_path = self.run_dir / f"{_safe_name(title, '课堂活动设计')}.docx"
        document = Document()
        document.styles["Normal"].font.name = "宋体"
        document.styles["Normal"].font.size = Pt(11)
        heading = document.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph(f"课程：{course_name}")
        for index, item in enumerate(activities, start=1):
            document.add_heading(f"活动{index}：{item.get('title') or '课堂活动'}", level=1)
            document.add_paragraph(f"目的：{item.get('purpose') or ''}")
            document.add_paragraph(
                f"形式：{item.get('format') or '小组'}　建议时间：{item.get('duration_minutes') or 0} 分钟"
            )
            document.add_heading("实施步骤", level=2)
            for step in item.get("instructions") or []:
                document.add_paragraph(str(step), style="List Number")
            document.add_heading("讨论问题", level=2)
            for question in item.get("questions") or []:
                document.add_paragraph(str(question), style="List Bullet")
            document.add_paragraph(
                f"资料依据：{'、'.join(item.get('evidence_refs') or []) or '请教师补充核验'}"
            )
            document.add_paragraph(f"评价标准：{item.get('evaluation') or ''}")
        self._append_sources(document, evidence)
        document.save(output_path)
        return {
            "kind": "classroom_activities",
            "title": title,
            "file_name": output_path.name,
            "storage_path": str(output_path.relative_to(self.root)),
            "media_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "activity_count": len(activities),
            "preview": activities,
        }

    @staticmethod
    def _append_sources(document: Document, evidence: list[dict[str, Any]]) -> None:
        document.add_heading("资料依据", level=1)
        for index, item in enumerate(evidence, start=1):
            title = item.get("source_title") or f"资料{index}"
            position = item.get("position") or ""
            document.add_paragraph(f"[资料{index}] {title}　{position}")

    def resolve_download(self, artifact: dict[str, Any]) -> Path:
        storage_path = artifact.get("storage_path")
        if not storage_path:
            raise FileNotFoundError("成果文件尚未生成")
        candidate = (self.root / str(storage_path)).resolve()
        root = self.root.resolve()
        if candidate != root and root not in candidate.parents:
            raise FileNotFoundError("成果文件路径无效")
        if not candidate.is_file():
            raise FileNotFoundError("成果文件不存在")
        return candidate
