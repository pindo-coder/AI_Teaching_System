from io import BytesIO
import logging
from urllib.parse import quote

from fastapi import APIRouter, Depends, Query
from fastapi.responses import Response
from sqlalchemy.orm import Session

from app.api.dependencies import get_current_user
from app.db.session import get_db
from app.models.user import User
from app.models.review_practice import ReviewPractice
from app.schemas.common import ApiResponse
from app.schemas.study import (
    NoteRelatedData, NoteSearchItem, ReviewAnswerResult, ReviewAnswerSubmit, ReviewQuestionRead, ReviewReferenceItem, ReviewReferencesRequest, ReviewSaveToNotesRequest, ReviewResultItem,
    ReviewRead, StudyChatHistorySave, StudyChatMessageRead, StudyNoteListItem, StudyNoteRead, StudyNoteUpdate,
    TextbookAnnotationCreate, TextbookAnnotationRead, TextbookAnnotationUpdate,
)
from app.services.study_service import StudyService
from app.services.task_service import TaskService
from app.schemas.task import LearningEventCreate


router = APIRouter(prefix="/study", tags=["study"])
logger = logging.getLogger(__name__)


@router.get("/chat-history/{chapter_id}", response_model=ApiResponse[list[StudyChatMessageRead]])
def chat_history(chapter_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)) -> ApiResponse[list[StudyChatMessageRead]]:
    return ApiResponse(data=StudyService(db).list_chat_history(user.id, chapter_id))


@router.post("/chat-history", response_model=ApiResponse[list[StudyChatMessageRead]])
def save_chat_history(payload: StudyChatHistorySave, user: User = Depends(get_current_user), db: Session = Depends(get_db)) -> ApiResponse[list[StudyChatMessageRead]]:
    return ApiResponse(message="AI 对话已保存", data=StudyService(db).save_chat_history(user.id, payload))


@router.delete("/chat-history/{chapter_id}", response_model=ApiResponse[dict[str, int]])
def clear_chat_history(chapter_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)) -> ApiResponse[dict[str, int]]:
    StudyService(db).clear_chat_history(user.id, chapter_id)
    return ApiResponse(message="本章会话已清空", data={"chapter_id": chapter_id})


@router.get("/notes", response_model=ApiResponse[list[StudyNoteListItem]])
def list_notes(user: User = Depends(get_current_user), db: Session = Depends(get_db)) -> ApiResponse[list[StudyNoteListItem]]:
    return ApiResponse(data=StudyService(db).list_notes(user.id))


@router.get("/notes/semantic-search", response_model=ApiResponse[list[NoteSearchItem]])
def search_notes(q: str = Query(min_length=1, max_length=200), course_id: int | None = None,
                 user: User = Depends(get_current_user), db: Session = Depends(get_db)) -> ApiResponse[list[NoteSearchItem]]:
    return ApiResponse(data=StudyService(db).search_notes(user.id, q, course_id))


@router.get("/notes/{chapter_id}/related", response_model=ApiResponse[NoteRelatedData])
def related_note_content(chapter_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)) -> ApiResponse[NoteRelatedData]:
    return ApiResponse(data=StudyService(db).related_note_content(user.id, chapter_id))


@router.get("/notes/{chapter_id}/export")
def export_note(chapter_id: int, format: str = Query(default="markdown", pattern="^(markdown|docx)$"),
                user: User = Depends(get_current_user), db: Session = Depends(get_db)) -> Response:
    content, title = StudyService(db).build_export_markdown(user.id, chapter_id)
    safe_title = quote(title.replace("/", "-"))
    if format == "markdown":
        return Response(content, media_type="text/markdown; charset=utf-8",
                        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{safe_title}-学习笔记.md"})
    from docx import Document
    document = Document()
    for line in content.splitlines():
        if line.startswith("# "):
            document.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            document.add_heading(line[4:], level=3)
        elif line:
            document.add_paragraph(line)
    buffer = BytesIO(); document.save(buffer)
    return Response(buffer.getvalue(), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": f"attachment; filename*=UTF-8''{safe_title}-学习笔记.docx"})


@router.get("/notes/{chapter_id}", response_model=ApiResponse[StudyNoteRead | None])
def get_note(chapter_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)) -> ApiResponse[StudyNoteRead | None]:
    return ApiResponse(data=StudyService(db).get_note(user.id, chapter_id))


@router.put("/notes/{chapter_id}", response_model=ApiResponse[StudyNoteRead])
def save_note(chapter_id: int, payload: StudyNoteUpdate, user: User = Depends(get_current_user), db: Session = Depends(get_db)) -> ApiResponse[StudyNoteRead]:
    service = StudyService(db)
    note = service.save_note(user.id, chapter_id, payload.content)
    try:
        TaskService(db).record(user.id, LearningEventCreate(
            course_id=note.course_id,
            chapter_id=note.chapter_id,
            learning_stage="review",
            event_type="note_saved",
            event_data={"content_length": len(service.plain_note_content(note.content))},
        ))
    except Exception:
        # The note has already been committed. Telemetry/task generation is
        # secondary and must not turn a successful note save into an error.
        db.rollback()
        logger.exception("study_note_event_record_failed user_id=%s chapter_id=%s", user.id, chapter_id)
    return ApiResponse(message="学习笔记已保存", data=note)


@router.delete("/notes/{note_id}", response_model=ApiResponse[dict[str, int]])
def delete_note(note_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)) -> ApiResponse[dict[str, int]]:
    StudyService(db).delete_note(user.id, note_id)
    return ApiResponse(message="学习笔记已删除", data={"id": note_id})


@router.get(
    "/textbook-annotations/chapters/{chapter_id}",
    response_model=ApiResponse[list[TextbookAnnotationRead]],
)
def list_textbook_annotations(
    chapter_id: int,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> ApiResponse[list[TextbookAnnotationRead]]:
    return ApiResponse(data=StudyService(db).list_textbook_annotations(user.id, chapter_id))


@router.post(
    "/textbook-annotations/chapters/{chapter_id}",
    response_model=ApiResponse[TextbookAnnotationRead],
    status_code=201,
)
def create_textbook_annotation(
    chapter_id: int,
    payload: TextbookAnnotationCreate,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> ApiResponse[TextbookAnnotationRead]:
    annotation = StudyService(db).create_textbook_annotation(user.id, chapter_id, payload)
    return ApiResponse(message="教材标注已保存", data=annotation)


@router.patch(
    "/textbook-annotations/{annotation_id}",
    response_model=ApiResponse[TextbookAnnotationRead],
)
def update_textbook_annotation(
    annotation_id: int,
    payload: TextbookAnnotationUpdate,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> ApiResponse[TextbookAnnotationRead]:
    annotation = StudyService(db).update_textbook_annotation(user.id, annotation_id, payload)
    return ApiResponse(message="教材标注已更新", data=annotation)


@router.delete("/textbook-annotations/{annotation_id}", response_model=ApiResponse[dict[str, int]])
def delete_textbook_annotation(
    annotation_id: int,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> ApiResponse[dict[str, int]]:
    StudyService(db).delete_textbook_annotation(user.id, annotation_id)
    return ApiResponse(message="教材标注已删除", data={"id": annotation_id})


@router.get("/reviews/today", response_model=ApiResponse[list[ReviewRead]])
def today_reviews(user: User = Depends(get_current_user), db: Session = Depends(get_db)) -> ApiResponse[list[ReviewRead]]:
    return ApiResponse(data=StudyService(db).due_reviews(user.id))


@router.post("/reviews/{chapter_id}/activate", response_model=ApiResponse[dict[str, int]])
def activate_review(chapter_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)) -> ApiResponse[dict[str, int]]:
    record = StudyService(db).activate_review(user.id, chapter_id)
    return ApiResponse(message="已加入间隔复习计划", data={"id": record.id, "interval_days": record.interval_days})


@router.post("/reviews/{chapter_id}/complete", response_model=ApiResponse[dict[str, int]])
def complete_review(chapter_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)) -> ApiResponse[dict[str, int]]:
    record = StudyService(db).complete_review(user.id, chapter_id)
    return ApiResponse(message="本次复习已完成", data={"id": record.id, "interval_days": record.interval_days})


@router.post("/reviews/{chapter_id}/questions", response_model=ApiResponse[list[ReviewQuestionRead]])
def review_questions(chapter_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)) -> ApiResponse[list[ReviewQuestionRead]]:
    records = StudyService(db).create_review_questions(user.id, chapter_id)
    return ApiResponse(message="已生成本章复习题", data=records)


@router.get("/reviews/{chapter_id}/latest-result", response_model=ApiResponse[list[ReviewResultItem]])
def latest_review_result(chapter_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)) -> ApiResponse[list[ReviewResultItem]]:
    return ApiResponse(data=StudyService(db).latest_review_result(user.id, chapter_id))


@router.post("/reviews/{chapter_id}/references", response_model=ApiResponse[list[ReviewReferenceItem]])
def review_references(chapter_id: int, payload: ReviewReferencesRequest, user: User = Depends(get_current_user), db: Session = Depends(get_db)) -> ApiResponse[list[ReviewReferenceItem]]:
    return ApiResponse(message="参考答案已生成", data=StudyService(db).review_references(user.id, chapter_id, payload.practice_ids, force=payload.force))


@router.post("/reviews/{chapter_id}/save-to-notes", response_model=ApiResponse[StudyNoteRead])
def save_review_to_notes(chapter_id: int, payload: ReviewSaveToNotesRequest, user: User = Depends(get_current_user), db: Session = Depends(get_db)) -> ApiResponse[StudyNoteRead]:
    note = StudyService(db).save_review_to_notes(user.id, chapter_id, payload.practice_ids)
    return ApiResponse(message="答题记录已保存到笔记空间", data=note)


@router.post("/reviews/questions/{practice_id}/answer", response_model=ApiResponse[ReviewAnswerResult])
def submit_review_answer(practice_id: int, payload: ReviewAnswerSubmit,
                         user: User = Depends(get_current_user), db: Session = Depends(get_db)) -> ApiResponse[ReviewAnswerResult]:
    practice = db.get(ReviewPractice, practice_id)
    result = StudyService(db).submit_review_answer(user.id, practice_id, payload.answer)
    if result.get("completed") and practice is not None and practice.user_id == user.id:
        TaskService(db).record(
            user.id,
            LearningEventCreate(
                course_id=practice.course_id,
                chapter_id=practice.chapter_id,
                learning_stage="exam",
                event_type="quiz_completed",
                event_data={"count": 1, "source": "spaced_review"},
            ),
        )
    return ApiResponse(message="作答已批阅", data=result)
