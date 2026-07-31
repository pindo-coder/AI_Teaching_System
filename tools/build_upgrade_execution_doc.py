from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "高校思政课AI智能教学辅助平台-版本改进与UI优化执行方案-V2.0-2026-07-27.docx"

FONT = "Arial Unicode MS"
INK = "202124"
MUTED = "5F6368"
BLUE = "1F4D78"
BLUE_LIGHT = "EAF2F8"
RED = "B71C1C"
RED_LIGHT = "FCE8E6"
GOLD = "B7791F"
GOLD_LIGHT = "FFF7E0"
GREEN = "137333"
GREEN_LIGHT = "E6F4EA"
GRAY_LIGHT = "F5F7F9"
GRAY_BORDER = "DADCE0"
WHITE = "FFFFFF"
CONTENT_DXA = 9360
TABLE_INDENT = 120


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=GRAY_BORDER, size=6):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths, indent=TABLE_INDENT):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_pr = table._tbl.tblPr
    tbl_w = table_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        table_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = table_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        table_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run(run, size=11, color=INK, bold=False, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_para(p, before=0, after=8, line=1.25, align=None, keep=False):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    if keep:
        p.paragraph_format.keep_with_next = True


def add_text(doc, text, *, size=11, color=INK, bold=False, italic=False,
             before=0, after=8, line=1.25, align=None, keep=False):
    p = doc.add_paragraph()
    set_para(p, before, after, line, align, keep)
    set_run(p.add_run(text), size, color, bold, italic)
    return p


def add_rich_para(doc, pieces, *, before=0, after=8, line=1.25, align=None, keep=False):
    p = doc.add_paragraph()
    set_para(p, before, after, line, align, keep)
    for text, opts in pieces:
        set_run(p.add_run(text), **opts)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_custom_numbering(doc, bullet=True):
    numbering = doc.part.numbering_part.element
    abs_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abs_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "280")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "290")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), str(abstract_id))
    num.append(abs_ref)
    numbering.append(num)
    return num_id


def add_list_item(doc, text, num_id, *, color=INK, bold_lead=None):
    p = doc.add_paragraph()
    set_para(p, 0, 4, 1.208)
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)
    if bold_lead and text.startswith(bold_lead):
        set_run(p.add_run(bold_lead), 11, color, True)
        set_run(p.add_run(text[len(bold_lead):]), 11, color)
    else:
        set_run(p.add_run(text), 11, color)
    return p


def add_table(doc, headers, rows, widths, *, header_fill=BLUE, font_size=9.5,
              first_col_bold=False, status_colors=False):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        set_para(p, 0, 0, 1.1, WD_ALIGN_PARAGRAPH.CENTER)
        set_run(p.add_run(header), font_size, WHITE, True)
    for row in rows:
        cells = table.add_row().cells
        for idx, (cell, value) in enumerate(zip(cells, row)):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            align = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 and widths[idx] < 1900 else WD_ALIGN_PARAGRAPH.LEFT
            set_para(p, 0, 0, 1.15, align)
            color = INK
            fill = None
            if status_colors and isinstance(value, str):
                if "通过" in value or "完成" in value:
                    color, fill = GREEN, GREEN_LIGHT
                elif "阻断" in value or "不得" in value:
                    color, fill = RED, RED_LIGHT
                elif "门槛" in value or "确认" in value:
                    color, fill = GOLD, GOLD_LIGHT
            set_run(p.add_run(str(value)), font_size, color, first_col_bold and idx == 0)
            if fill:
                set_cell_shading(cell, fill)
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)
            set_cell_margins(cell)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc, label, text, fill=BLUE_LIGHT, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, accent, 8)
    p = cell.paragraphs[0]
    set_para(p, 0, 0, 1.2)
    set_run(p.add_run(f"{label}  "), 10.5, accent, True)
    set_run(p.add_run(text), 10.5, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run(run, 9, MUTED)


def setup_styles(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.42)
    sec.footer_distance = Inches(0.42)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    settings = [
        ("Title", 30, BLUE, 0, 8),
        ("Subtitle", 14, MUTED, 0, 18),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, RED, 8, 4),
    ]
    for name, size, color, before, after in settings:
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name.startswith("Heading")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for section in doc.sections:
        hp = section.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_run(hp.add_run("高校思政课 AI 智能教学辅助平台  ·  版本改进与 UI 优化执行方案"), 8.5, MUTED)
        fp = section.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_run(fp.add_run("V2.0  |  2026-07-27  |  "), 8.5, MUTED)
        add_page_number(fp)


def add_cover(doc):
    add_text(doc, "PRODUCT UPGRADE & UI EXECUTION", size=10, color=RED, bold=True,
             before=26, after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("高校思政课 AI 智能教学辅助平台")
    p2 = doc.add_paragraph()
    set_para(p2, 0, 8, 1.1, WD_ALIGN_PARAGRAPH.CENTER)
    set_run(p2.add_run("版本改进与 UI 优化执行方案"), 24, RED, True)
    add_text(doc, "从“功能堆叠”升级为权威资料驱动、教师可控、学生可参与的教学闭环",
             size=13, color=MUTED, after=26, line=1.2, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_callout(
        doc,
        "执行主线",
        "V0 体验与信息架构冻结 → V1 竞赛闭环 → V2 教学班试点 → V3 规模化与多模态",
        GOLD_LIGHT,
        GOLD,
    )
    add_text(doc, "适用范围", size=10, color=MUTED, bold=True, before=12, after=4)
    add_text(doc, "竞赛答辩、教师试用、单教学班试点及后续产品化建设",
             size=11, color=INK, after=20)
    add_text(doc, "版本 V2.0  ·  2026年7月27日", size=10.5, color=MUTED,
             after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "基于《重大升级方案（2026-07-24）》与现有 UI.pdf 审阅重构",
             size=9.5, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


def add_exec_summary(doc, bullet_id):
    add_heading(doc, "0. 执行摘要", 1)
    add_callout(
        doc,
        "核心决策",
        "先冻结角色化信息架构和 UI 设计系统，再围绕“权威更新、智能备课、师生共建”三条闭环分版本交付；任何新增功能必须落入明确角色、课程上下文和可验收任务。",
        RED_LIGHT,
        RED,
    )
    add_heading(doc, "0.1 本次优化解决什么", 2)
    for text in [
        "方案层：原方案功能定义完整，但 14 个章节平铺，缺少版本门、依赖关系、淘汰边界和按周可验收结果。",
        "产品层：教师、学生、管理员共享大量并列入口，用户必须先理解系统模块，才能完成真实教学任务。",
        "UI 层：现稿大面积高饱和红色、装饰性背景和巨型卡片竞争注意力，层级、密度和状态语义不足。",
        "工程层：资料发现、Agent、通知与多模态都依赖后台任务、权限、审计和对象存储，需要先建设共同底座。",
    ]:
        add_list_item(doc, text, bullet_id)
    add_heading(doc, "0.2 优化后的产品承诺", 2)
    add_table(
        doc,
        ["对象", "进入平台后第一件事", "平台给出的确定价值"],
        [
            ("教师", "看到今天的备课、资料变化与待审核事项", "在有依据、可编辑、可追溯的前提下完成备课与发布"),
            ("学生", "看到今日学习、待完成与教师反馈", "把教材、讨论、笔记和复习串成连续学习路径"),
            ("管理员", "看到来源异常、候选资料与内容风险", "确保资料权威、权限可控、变更可审计"),
        ],
        [1500, 3300, 4560],
        header_fill=BLUE,
        first_col_bold=True,
    )
    add_heading(doc, "0.3 成功判断", 2)
    add_rich_para(
        doc,
        [
            ("不是“上线更多 AI 按钮”，而是 ", {"size": 11, "color": INK}),
            ("教师能在 5 步内完成一次有依据的备课、学生能在 2 步内进入今日任务、管理员能在 1 个审核页完成资料确认。", {"size": 11, "color": RED, "bold": True}),
        ],
        after=10,
    )


def add_phase_plan(doc, bullet_id):
    add_heading(doc, "1. 分阶段版本路线", 1)
    add_text(doc, "阶段划分采用“版本门”而非功能清单：上一阶段未通过验收，不进入下一阶段扩展。",
             color=MUTED, italic=True)
    rows = [
        ("V0｜体验基线", "3–5天", "角色任务、导航、UI令牌、页面模板、权限矩阵", "三类角色首要任务均有唯一入口；桌面一级导航≤5项"),
        ("V1｜竞赛闭环", "2周", "权威更新、智能备课、课堂发布三条可演示链路", "端到端脚本可重复；关键证据可核验；所有发布需确认"),
        ("V2｜教学班试点", "3–4周", "真实数据、共建审核、通知、PWA、运行监控", "1个教学班稳定使用；无跨班泄漏；任务与通知可追踪"),
        ("V3｜规模化", "4–6周", "多模态、模板中心、性能与成本治理、学校级配置", "50人在线稳定；AI超限排队；成本、失败和保留期可观测"),
    ]
    add_table(doc, ["版本门", "周期", "交付范围", "出门条件"], rows,
              [1650, 1050, 3150, 3510], header_fill=RED, first_col_bold=True)

    add_heading(doc, "1.1 V0：体验基线冻结", 2)
    for text in [
        "角色路径：分别完成教师“备课—发布”、学生“学习—反馈”、管理员“发现—审核”任务流。",
        "信息架构：冻结一级导航、课程上下文、全局消息入口与长任务入口。",
        "设计系统：冻结颜色、字号、间距、圆角、阴影、状态和按钮语义；不再以页面单独定风格。",
        "原型门：四个关键页面完成中保真走查并通过可用性评审，再进入开发。",
    ]:
        add_list_item(doc, text, bullet_id)

    add_heading(doc, "1.2 V1：竞赛与答辩版本", 2)
    add_table(
        doc,
        ["闭环", "演示脚本", "必须真实", "可暂用样例数据"],
        [
            ("权威更新", "发现新材料→显示新旧表述→管理员确认→提醒教师", "来源、原文、时间、审核记录", "检索任务可预置1条成功案例"),
            ("智能备课", "选专题/课时→证据包→课纲→PPT大纲→发布讨论", "引用、可编辑成果、发布确认", "PPT模板可先固定1套"),
            ("课堂衔接", "教师发布→学生进入任务→提交观点→教师反馈", "班级权限、状态流转、反馈记录", "学生样例账号与演示投稿"),
        ],
        [1600, 3300, 2500, 1960],
        header_fill=BLUE,
        first_col_bold=True,
    )
    add_heading(doc, "1.3 V2：单教学班试点", 2)
    for text in [
        "接入 3–5 个稳定权威来源，资料候选池开始按日运行并记录失败原因。",
        "上线专题讨论、学习卡片、AI初审和教师复核，积分仅针对教师采用或有效贡献。",
        "上线站内通知与 PWA，紧急/重要/普通/观察四级规则真实生效。",
        "建立错误率、任务队列、模型成本、通知送达与权限审计看板。",
    ]:
        add_list_item(doc, text, bullet_id)
    add_heading(doc, "1.4 V3：规模化与多模态", 2)
    for text in [
        "图片 OCR、教材拍照问答和课堂录音转写按保留期、权限和删除策略上线。",
        "学校模板、教师个人模板与成果版本中心形成可维护体系。",
        "数据库、任务 Worker、对象存储和 API 服务按负载拆分；对模型限流和排队透明展示。",
    ]:
        add_list_item(doc, text, bullet_id)


def add_scope(doc):
    add_heading(doc, "2. 优先级与范围边界", 1)
    add_table(
        doc,
        ["现在做（V0–V1）", "下一步（V2）", "稍后做（V3）", "明确不做"],
        [
            ("角色化导航、教师工作台", "学生共建与审核", "图片OCR/图片问答", "未经确认直接发布中央材料"),
            ("资料候选与差异证据", "通知中心/PWA", "课堂录音转写", "自动群发教学任务或全平台通知"),
            ("备课Agent最小闭环", "真实来源调度与监控", "模板中心/学校级配置", "人脸、声纹和身份推断"),
            ("成果编辑/导出/发布确认", "单班运行数据与审计", "50人并发与成本治理", "竞赛阶段本地部署大模型/GPU"),
        ],
        [2340, 2340, 2340, 2340],
        header_fill=BLUE,
    )
    add_callout(
        doc,
        "范围纪律",
        "新增需求如果不能直接提升三条闭环的完成率、证据可信度或教师控制力，就进入后续版本池，不插入当前迭代。",
        GOLD_LIGHT,
        GOLD,
    )


def add_ui_audit(doc):
    add_heading(doc, "3. 现有 UI 审阅与改进方向", 1)
    add_heading(doc, "3.1 已识别的主要问题", 2)
    add_table(
        doc,
        ["问题", "现稿表现", "带来的影响", "优化动作"],
        [
            ("视觉权重失衡", "大面积高饱和红底、黄边、纹样与巨型标题连续出现", "所有模块都像重点，阅读疲劳，状态难区分", "白/浅灰为主表面；红仅用于权威/警示；蓝用于操作"),
            ("首页任务不聚焦", "欢迎横幅、任务、进度、路径、入口卡片重复表达", "首屏不清楚今天最该做什么", "首屏固定“今日待办+变化+最近成果”，其他下沉"),
            ("角色路径混合", "学生式课程学习与教师备课、资料审核缺少明确分层", "不同角色看到无关入口", "登录后按角色加载独立导航与工作台模板"),
            ("操作语义不统一", "红、黄、绿、蓝按钮并存，主次与风险不稳定", "用户无法预判操作后果", "一个页面一个主按钮；发布/删除有统一确认样式"),
            ("信息密度不稳定", "有的区域过满，有的编辑器和侧栏大片留白", "难以扫描，空间利用低", "采用12列栅格、固定侧栏宽度与空状态组件"),
            ("交互状态缺失", "长任务、刷新、AI总结、生成笔记缺少完整状态", "失败、等待、可取消边界不清", "统一排队/执行/待确认/成功/失败/取消状态机"),
        ],
        [1550, 2700, 2300, 2810],
        header_fill=RED,
        font_size=8.8,
        first_col_bold=True,
    )
    add_heading(doc, "3.2 保留的设计资产", 2)
    add_text(doc, "保留并收敛，而不是全部推翻：", color=MUTED)
    add_table(
        doc,
        ["保留", "调整方式"],
        [
            ("中国红与传统文化纹样", "仅用于品牌页、权威资料标签和重点横幅；正文区降低饱和度与占比"),
            ("课程—时政—互动三位一体叙事", "转为真实任务路径和数据状态，而不是三张宣传卡片"),
            ("专题笔记三栏结构", "保留索引/编辑/AI辅助，但增加折叠、引用、版本与空状态"),
            ("现有长城、华表、城市线稿", "作为低对比装饰资产，不压在正文和控件背后"),
        ],
        [2600, 6760],
        header_fill=BLUE,
        first_col_bold=True,
    )


def add_ia(doc):
    add_heading(doc, "4. 角色化信息架构", 1)
    add_heading(doc, "4.1 一级导航（桌面端≤5项）", 2)
    add_table(
        doc,
        ["角色", "一级导航", "工作台首屏", "全局固定入口"],
        [
            ("教师", "工作台｜课程备课｜课堂教学｜师生共建｜资料动态", "待备课、资料变化、待审核、最近成果", "消息、长任务、全局搜索、个人设置"),
            ("学生", "今日学习｜课程｜讨论共建｜笔记复习｜消息", "待完成、当前专题、教师反馈、今日复习", "课程切换、搜索、个人成长"),
            ("管理员", "平台概览｜教学管理｜资料审核｜用户权限｜系统设置", "来源异常、候选资料、内容风险、系统状态", "消息、后台任务、审计、帮助"),
        ],
        [1200, 3500, 3100, 1560],
        header_fill=BLUE,
        font_size=9,
        first_col_bold=True,
    )
    add_heading(doc, "4.2 页面上下文规则", 2)
    add_callout(
        doc,
        "唯一上下文",
        "任何问答、生成、通知与共建内容必须绑定“角色 + 课程 + 专题/章节 + 教学班 + 资料范围”；上下文显示在页面顶部，可切换但不可隐式漂移。",
        BLUE_LIGHT,
        BLUE,
    )
    add_table(
        doc,
        ["规则", "实现要求"],
        [
            ("唯一主入口", "同一业务只保留一个主入口；其他页面用上下文快捷入口跳转"),
            ("唯一主操作", "每页最多一个实心主按钮；次要操作用描边或文字；危险操作固定红色"),
            ("逐步展开", "默认展示任务必需信息；模型、检索范围、模板等放入高级设置"),
            ("深度链接", "通知必须跳到具体材料、任务、讨论或成果，不落到泛化首页"),
        ],
        [1900, 7460],
        header_fill=RED,
        first_col_bold=True,
    )


def add_design_system(doc):
    add_heading(doc, "5. UI 设计系统（V0 冻结项）", 1)
    add_heading(doc, "5.1 颜色与语义", 2)
    add_table(
        doc,
        ["令牌", "建议值", "使用范围", "禁止"],
        [
            ("Surface/页面", "#F7F8FA / #FFFFFF", "页面背景、卡片、编辑区", "正文大面积红底"),
            ("Authority/权威红", "#B71C1C", "中央资料、重大提醒、风险确认", "普通主按钮和所有标题通用"),
            ("Action/科技蓝", "#1F4D78", "主按钮、链接、选中态、进度", "与红色同时争夺主操作"),
            ("Highlight/精选金", "#B7791F", "精选、徽章、少量重点", "正文、长段文字和大面积背景"),
            ("Success/通过", "#137333", "审核通过、完成状态", "非状态性装饰"),
            ("Text/正文", "#202124 / #5F6368", "正文、次级说明", "浅灰小字叠加复杂背景"),
        ],
        [1800, 1600, 3100, 2860],
        header_fill=BLUE,
        font_size=9,
        first_col_bold=True,
    )
    add_heading(doc, "5.2 尺寸令牌", 2)
    add_table(
        doc,
        ["项目", "桌面端", "移动端", "规则"],
        [
            ("内容宽度", "1200px，12列栅格", "单列，左右16px", "不让正文跨越过宽"),
            ("顶部栏/侧栏", "64px / 224px", "56px / 抽屉", "侧栏可折叠，保留图标+文字"),
            ("字体", "32/24/20/16/14px", "28/22/18/16/14px", "正文最小14px；表格不小于13px"),
            ("间距", "4/8/12/16/24/32/48", "4/8/12/16/24/32", "只使用8px基线体系"),
            ("圆角", "卡片12px，控件8px", "卡片12px，控件8px", "避免所有元素胶囊化"),
            ("阴影", "0 2px 8px rgba(0,0,0,.08)", "同桌面", "仅浮层/悬浮卡片；普通卡片用边框"),
        ],
        [1500, 2100, 2100, 3660],
        header_fill=RED,
        font_size=9,
        first_col_bold=True,
    )
    add_heading(doc, "5.3 核心组件清单", 2)
    add_text(
        doc,
        "AppShell、RoleNav、CourseContext、TaskCard、EvidenceCard、StatusChip、SourceBadge、StepProgress、DiffViewer、ArtifactEditor、SmartPanel、ConfirmDialog、NotificationItem、EmptyState、Skeleton、ErrorState。",
        color=INK,
    )


def add_wireframes(doc):
    add_heading(doc, "6. 关键页面 UI 设计", 1)
    add_text(doc, "以下为开发可执行的结构线框。视觉稿应遵循同一栅格、颜色与组件令牌，不再逐页另起样式。",
             color=MUTED, italic=True)

    add_heading(doc, "6.1 教师工作台", 2)
    add_table(
        doc,
        ["左侧导航 224px", "主工作区（8列）", "右侧辅助区（4列）"],
        [
            ("工作台【选中】\n课程备课\n课堂教学\n师生共建\n资料动态",
             "顶部：课程/教学班上下文\n\n今日待办（按截止和风险排序）\n· 待备课 1\n· 待发布 2\n· 待审核 3\n\n最近成果：课纲/PPT/讨论",
             "资料变化\n· 1条新表述待确认\n· 2条相关材料\n\n快速操作\n【开始智能备课】\n查看长任务"),
        ],
        [1900, 4700, 2760],
        header_fill=BLUE,
        font_size=9.3,
    )
    add_callout(doc, "首屏原则", "只回答“今天要做什么、发生了什么变化、从哪里继续”，不展示平台全部能力。", BLUE_LIGHT, BLUE)

    add_heading(doc, "6.2 权威资料候选审核", 2)
    add_table(
        doc,
        ["候选列表", "证据与差异主区", "审核面板"],
        [
            ("筛选：来源/时间/相关度/状态\n\n候选卡\n标题 + 来源徽标\n发布时间 + 相关专题\n新表述标签",
             "原文信息与网页快照\n\n【教材/旧材料】 ⇄ 【新材料】\n段落级差异，高亮新增/删除\n\nAI判断：相关性、权威性、时效性\n每个判断都可回到原文",
             "权威层级\n适用课程/专题\n提醒级别\n有效期\n\n【退回候选】\n【确认并发布】"),
        ],
        [2100, 4700, 2560],
        header_fill=RED,
        font_size=8.8,
    )
    add_callout(doc, "确认边界", "AI可检索、摘要、关联和比较；正式发布、全平台提醒和教材解释调整必须人工确认。", RED_LIGHT, RED)

    add_heading(doc, "6.3 智能备课 Agent", 2)
    add_table(
        doc,
        ["步骤/资料", "成果编辑区", "AI辅助与引用"],
        [
            ("1 设置任务  ✓\n2 构建证据  ✓\n3 生成课纲  ●\n4 生成成果\n5 预览发布\n\n资料范围\n教材 12段\n中央材料 4条\n教师资料 2份",
             "标题：第×专题教学设计\n\n目标｜重点难点｜教学过程｜活动｜评价\n支持局部重写、版本比较与自动保存\n\n底部固定：\n【保存草稿】 【下一步：生成PPT大纲】",
             "引用检查\n· 3处已校准\n· 1处待确认\n\n局部建议\n· 加入新表述对照\n· 生成讨论题\n\n长任务状态/取消/重试"),
        ],
        [2100, 4700, 2560],
        header_fill=BLUE,
        font_size=8.8,
    )
    add_heading(doc, "6.4 学生今日学习（移动优先）", 2)
    add_table(
        doc,
        ["首屏顺序", "组件", "交互要求"],
        [
            ("1", "今日待办", "显示任务、截止、预计时长；首个任务一键进入"),
            ("2", "继续学习", "恢复上次专题、阅读位置和未完成步骤"),
            ("3", "教师反馈", "跳到具体讨论/卡片，不跳泛化消息页"),
            ("4", "今日复习", "最多3条，支持完成、稍后和解释原因"),
            ("5", "重要通知", "仅显示紧急/重要；普通信息进入消息中心"),
        ],
        [900, 2600, 5860],
        header_fill=RED,
        first_col_bold=True,
    )


def add_interactions(doc, bullet_id):
    add_heading(doc, "7. 关键交互与状态规范", 1)
    add_heading(doc, "7.1 长任务状态机", 2)
    add_table(
        doc,
        ["状态", "界面反馈", "允许操作", "通知"],
        [
            ("排队中", "显示队列位置/预计等待", "取消、后台运行", "不通知"),
            ("执行中", "步骤、进度、当前资料与耗时", "暂停、取消、查看日志摘要", "离页后完成再通知"),
            ("待确认", "突出需用户决策的差异/风险", "确认、替换资料、返回上一步", "重要通知"),
            ("成功", "给出成果链接与下一步", "编辑、导出、发布", "普通通知"),
            ("失败", "人类可读原因 + 失败步骤", "重试、换资料、下载日志编号", "重要通知"),
            ("已取消", "保留已完成步骤和草稿", "重新开始、删除草稿", "不通知"),
        ],
        [1250, 3300, 3000, 1810],
        header_fill=BLUE,
        font_size=9,
        first_col_bold=True,
    )
    add_heading(doc, "7.2 高风险操作", 2)
    for text in [
        "发布教学任务：确认教学班、可见时间、截止时间和通知对象。",
        "确认中央材料：确认权威层级、适用范围、有效期、原文快照和提醒级别。",
        "群发通知：预览接收人数、深度链接和去重键；不可跨班级默认发送。",
        "删除媒体或成果：说明影响范围和是否可恢复；原始音频按保留期自动清理。",
    ]:
        add_list_item(doc, text, bullet_id)
    add_heading(doc, "7.3 空状态、失败与可访问性", 2)
    for text in [
        "空状态必须解释为什么为空，并提供唯一下一步；不得用大片空白或无意义插画占位。",
        "所有状态不只依赖颜色，必须同时有文字/图标；交互控件具备键盘焦点与可读标签。",
        "正文与背景对比度满足 WCAG AA；移动端触控目标不小于 44×44px。",
    ]:
        add_list_item(doc, text, bullet_id)


def add_engineering(doc):
    add_heading(doc, "8. 产品—工程落地映射", 1)
    add_table(
        doc,
        ["前端能力", "核心实体/服务", "后台任务", "审计/权限"],
        [
            ("角色工作台", "Course、Class、Task、Artifact", "聚合待办与变化", "按角色/教学班过滤"),
            ("资料候选审核", "SourceRegistry、DiscoveryJob、MaterialCandidate、PolicyChange", "检索、抓取、解析、去重、差异", "确认人、原文快照、内容哈希"),
            ("智能备课", "AgentRun、AgentStep、GeneratedArtifact、ArtifactTemplate", "检索证据、生成、导出PPT", "步骤输入/模型/输出/版本"),
            ("共建与审核", "CommunityPost、Reply、LearningCard、Review", "AI初审、重复检测", "作者、可见范围、审核链"),
            ("通知/PWA", "Notification、Subscription、PushToken", "批量发送、去重、重试", "对象、级别、深度链接、已读"),
            ("图片/音频", "MediaAsset、Transcript", "OCR、转写、清理", "保留期、删除、班级权限"),
        ],
        [2100, 3100, 2200, 1960],
        header_fill=RED,
        font_size=8.5,
        first_col_bold=True,
    )
    add_heading(doc, "8.1 共同技术底座（V1 前完成）", 2)
    add_table(
        doc,
        ["底座", "最低实现", "阻断条件"],
        [
            ("后台队列", "Redis + Worker；幂等、重试、超时、取消、结果链接", "长任务仍阻塞 API"),
            ("对象存储", "教材、快照、PPT、图片、音频统一元数据与权限", "媒体继续堆在本地系统盘"),
            ("权限与审计", "高风险操作、跨班访问、发布/删除均记录", "无法追溯操作人和影响范围"),
            ("模型适配层", "文本/Embedding/视觉/语音统一接口与配额", "业务代码绑定单一模型厂商"),
            ("可观测性", "任务失败率、耗时、队列、成本、送达率", "试点期间无法定位失败"),
        ],
        [1800, 4960, 2600],
        header_fill=BLUE,
        first_col_bold=True,
        status_colors=True,
    )


def add_acceptance(doc):
    add_heading(doc, "9. 验收指标与评审门", 1)
    add_table(
        doc,
        ["领域", "V1目标", "V2目标", "验收方式"],
        [
            ("可用性", "教师备课≤5步；学生待办≤2步", "核心任务成功率≥85%", "5名目标用户任务走查"),
            ("资料可信", "关键结论100%可回到来源", "来源成功率≥95%，重复率≤5%", "抽样核验+任务日志"),
            ("生成成果", "课纲/PPT大纲可编辑可追溯", "生成/导出成功率≥95%", "端到端回归"),
            ("权限合规", "发布/通知/删除均确认", "无跨班泄漏；审计完整", "权限矩阵与渗透用例"),
            ("性能", "演示链路稳定，失败可恢复", "50人在线无OOM；AI超限排队", "压测与故障演练"),
            ("UI一致性", "关键页使用统一令牌/组件", "无页面级颜色和按钮语义漂移", "设计系统审计"),
        ],
        [1500, 2550, 2700, 2610],
        header_fill=RED,
        font_size=8.8,
        first_col_bold=True,
    )
    add_heading(doc, "9.1 评审节奏", 2)
    add_table(
        doc,
        ["时间", "评审内容", "参与者", "输出"],
        [
            ("每周一", "范围、依赖、风险与本周验收项", "产品/教学/设计/研发", "迭代清单与阻断项"),
            ("每周三", "关键页面与端到端链路走查", "设计/前端/后端/测试", "问题单与决策记录"),
            ("每周五", "可演示增量与数据复盘", "项目组+教师代表", "通过/不通过与下周入口条件"),
        ],
        [1300, 3500, 2400, 2160],
        header_fill=BLUE,
        first_col_bold=True,
    )


def add_backlog(doc):
    add_heading(doc, "10. 首个两周执行清单", 1)
    add_table(
        doc,
        ["日程", "产品/设计", "前端", "后端/AI", "验收结果"],
        [
            ("D1–D2", "冻结角色路径、导航、UI令牌", "搭建AppShell与路由骨架", "冻结权限矩阵与任务状态机", "V0评审通过"),
            ("D3–D4", "教师工作台/资料审核中保真", "TaskCard、EvidenceCard、状态组件", "候选/差异接口与样例数据", "关键页可点击"),
            ("D5", "备课Agent线框与文案", "StepProgress与编辑器框架", "AgentRun/Step与证据包", "端到端空链路"),
            ("D6–D7", "发布确认/失败态/空状态", "资料差异与成果编辑", "检索、生成、任务队列", "三条闭环联调"),
            ("D8–D9", "答辩脚本、提示文案、演示数据", "响应式/加载/错误处理", "导出、通知、审计", "故障场景可恢复"),
            ("D10", "统一性审计与可用性走查", "修复与构建", "数据备份与监控", "V1发布候选"),
        ],
        [1150, 2350, 2100, 2250, 1510],
        header_fill=RED,
        font_size=8.2,
        first_col_bold=True,
    )
    add_callout(
        doc,
        "立即执行顺序",
        "先完成 D1–D5 的体验基线与空链路；只有关键页面、状态机和数据契约冻结后，才进入真实抓取、生成与通知联调。",
        GOLD_LIGHT,
        GOLD,
    )


def add_decisions(doc):
    add_heading(doc, "附录A：关键决策记录", 1)
    add_table(
        doc,
        ["议题", "V2.0确认决策"],
        [
            ("产品定位", "权威资料驱动的高校思政课程智能教学平台，不是通用聊天机器人"),
            ("版本策略", "四个版本门；V1只保留三条可演示闭环"),
            ("UI策略", "中性表面为主，权威红/操作蓝语义分离，角色化工作台"),
            ("中央资料", "自动发现进入候选池；正式发布、重大提醒与教材解释调整人工确认"),
            ("Agent权限", "生成/修改草稿可自动；发布、群发、删除必须确认"),
            ("PPT", "模型输出页面级结构，服务端按模板生成可编辑文件"),
            ("共建", "AI初审+教师复核；个人成长为主，不默认公开排名"),
            ("多模态", "图片和音频优先；原始音频按保留期清理；不做人脸/声纹"),
            ("基础设施", "V2试点建议4核8GB、10Mbps、Redis Worker与对象存储"),
        ],
        [2200, 7160],
        header_fill=BLUE,
        first_col_bold=True,
    )
    add_heading(doc, "附录B：交付物清单", 1)
    add_table(
        doc,
        ["阶段", "文档/设计", "软件", "数据与运维"],
        [
            ("V0", "角色任务流、IA、UI令牌、4页中保真、组件清单", "AppShell/路由/状态组件骨架", "权限矩阵、状态机、接口契约"),
            ("V1", "答辩脚本、演示文案、异常态", "三条闭环、成果编辑、发布确认", "样例数据、审计、备份、演示监控"),
            ("V2", "学生/管理员全量页面、PWA规范", "共建、通知、真实调度", "运行看板、压测、故障演练"),
            ("V3", "多模态/模板中心规范", "OCR、转写、学校级配置", "成本、保留期、扩容方案"),
        ],
        [1200, 3200, 2600, 2360],
        header_fill=RED,
        first_col_bold=True,
    )


def main():
    doc = Document()
    setup_styles(doc)
    bullet_id = add_custom_numbering(doc, bullet=True)
    add_custom_numbering(doc, bullet=False)
    add_cover(doc)
    add_exec_summary(doc, bullet_id)
    add_phase_plan(doc, bullet_id)
    add_scope(doc)
    add_ui_audit(doc)
    add_ia(doc)
    add_design_system(doc)
    add_wireframes(doc)
    add_interactions(doc, bullet_id)
    add_engineering(doc)
    add_acceptance(doc)
    add_backlog(doc)
    add_decisions(doc)

    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.8)
        section.header_distance = Inches(0.42)
        section.footer_distance = Inches(0.42)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
