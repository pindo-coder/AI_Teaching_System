from __future__ import annotations

import hashlib
import math
import shutil
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs" / "软件著作权申请材料"
ASSET_DIR = OUTPUT_DIR / ".build_assets"

BODY_ASCII = "Hiragino Sans GB"
BODY_CJK = "Hiragino Sans GB"
CODE_ASCII = "Menlo"
CODE_CJK = "Hiragino Sans GB"

INK = "202124"
MUTED = "5F6368"
NAVY = "1F4D78"
BLUE = "2E74B5"
RED = "A61B1B"
GOLD = "8A5A00"
GREEN = "176B4D"
PALE_BLUE = "E8EEF5"
PALE_RED = "FCE8E6"
PALE_GOLD = "FFF4D6"
PALE_GREEN = "E8F3EE"
PALE_GRAY = "F2F4F7"
BORDER = "DADCE0"
WHITE = "FFFFFF"

CONTENT_DXA = 9360
TABLE_INDENT = 120


@dataclass(frozen=True)
class CodeExcerpt:
    title: str
    purpose: str
    source: str
    ranges: tuple[tuple[int, int], ...]


@dataclass(frozen=True)
class SoftwareSpec:
    output_name: str
    software_name: str
    short_name: str
    subtitle: str
    scope_statement: str
    running_label: str
    diagram_name: str
    diagram_steps: tuple[str, ...]
    sections: tuple[tuple[str, tuple[str, ...]], ...]
    code_excerpts: tuple[CodeExcerpt, ...]


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run_font(
    run,
    *,
    size: float = 10.5,
    color: str = INK,
    bold: bool = False,
    italic: bool = False,
    ascii_font: str = BODY_ASCII,
    cjk_font: str = BODY_CJK,
) -> None:
    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), ascii_font)
    fonts.set(qn("w:hAnsi"), ascii_font)
    fonts.set(qn("w:eastAsia"), cjk_font)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic


def set_paragraph(
    paragraph,
    *,
    before: float = 0,
    after: float = 6,
    line: float = 1.10,
    align=None,
    keep_with_next: bool = False,
) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep_with_next
    if align is not None:
        paragraph.alignment = align


def shade_paragraph(paragraph, fill: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_shading(cell, fill: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    margins = tcpr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tcpr.append(margins)
    for name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width: int) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    tcw = tcpr.find(qn("w:tcW"))
    if tcw is None:
        tcw = OxmlElement("w:tcW")
        tcpr.append(tcw)
    tcw.set(qn("w:w"), str(width))
    tcw.set(qn("w:type"), "dxa")


def set_cell_borders(cell, color: str = BORDER, size: int = 6) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    borders = tcpr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcpr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths: Iterable[int], indent: int = TABLE_INDENT) -> None:
    widths = list(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblpr = table._tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(sum(widths)))
    tblw.set(qn("w:type"), "dxa")
    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), str(indent))
    tblind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        trpr = row._tr.get_or_add_trPr()
        if trpr.find(qn("w:cantSplit")) is None:
            trpr.append(OxmlElement("w:cantSplit"))
        for cell, width in zip(row.cells, widths, strict=True):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))
    set_run_font(run, size=8.5, color=MUTED)


def add_text(
    doc: Document,
    text: str,
    *,
    size: float = 10.5,
    color: str = INK,
    bold: bool = False,
    italic: bool = False,
    before: float = 0,
    after: float = 6,
    line: float = 1.10,
    align=None,
    keep_with_next: bool = False,
):
    paragraph = doc.add_paragraph()
    set_paragraph(
        paragraph,
        before=before,
        after=after,
        line=line,
        align=align,
        keep_with_next=keep_with_next,
    )
    set_run_font(paragraph.add_run(text), size=size, color=color, bold=bold, italic=italic)
    return paragraph


def add_rich_text(doc: Document, parts: list[tuple[str, dict]], **paragraph_options):
    paragraph = doc.add_paragraph()
    set_paragraph(paragraph, **paragraph_options)
    for text, options in parts:
        set_run_font(paragraph.add_run(text), **options)
    return paragraph


def add_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)
    return paragraph


def add_numbering(doc: Document, *, bullet: bool) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(item.get(qn("w:abstractNumId"))) for item in numbering.findall(qn("w:abstractNum"))]
    number_ids = [int(item.get(qn("w:numId"))) for item in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    number_id = max(number_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    numfmt = OxmlElement("w:numFmt")
    numfmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    level.append(numfmt)
    leveltext = OxmlElement("w:lvlText")
    leveltext.set(qn("w:val"), "●" if bullet else "%1.")
    level.append(leveltext)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    level.append(suffix)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ppr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    ppr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)
    level.append(ppr)
    abstract.append(level)
    numbering.append(abstract)

    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(number_id))
    reference = OxmlElement("w:abstractNumId")
    reference.set(qn("w:val"), str(abstract_id))
    number.append(reference)
    numbering.append(number)
    return number_id


def add_list_item(doc: Document, text: str, number_id: int, *, bold_lead: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph(paragraph, after=8, line=1.167)
    ppr = paragraph._p.get_or_add_pPr()
    numpr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(number_id))
    numpr.extend((ilvl, numid))
    ppr.append(numpr)
    if bold_lead and text.startswith(bold_lead):
        set_run_font(paragraph.add_run(bold_lead), bold=True)
        set_run_font(paragraph.add_run(text[len(bold_lead):]))
    else:
        set_run_font(paragraph.add_run(text))


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[tuple],
    widths: list[int],
    *,
    header_fill: str = NAVY,
    first_col_bold: bool = False,
    font_size: float = 9.2,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for cell, width in zip(table.rows[0].cells, widths, strict=True):
        set_cell_width(cell, width)
    set_table_geometry(table, widths)
    header_row = table.rows[0]
    header_row._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for cell, header in zip(header_row.cells, headers, strict=True):
        set_cell_shading(cell, header_fill)
        paragraph = cell.paragraphs[0]
        set_paragraph(paragraph, after=0, line=1.05, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_run_font(paragraph.add_run(header), size=font_size, color=WHITE, bold=True)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, (cell, value) in enumerate(zip(cells, values, strict=True)):
            if row_index % 2:
                set_cell_shading(cell, "F8F9FA")
            paragraph = cell.paragraphs[0]
            align = WD_ALIGN_PARAGRAPH.CENTER if widths[index] <= 1600 else WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph(paragraph, after=0, line=1.10, align=align)
            set_run_font(
                paragraph.add_run(str(value)),
                size=font_size,
                bold=first_col_bold and index == 0,
            )
    set_table_geometry(table, widths)
    add_text(doc, "", after=2, size=2)


def add_callout(doc: Document, label: str, text: str, *, fill: str = PALE_BLUE, accent: str = NAVY) -> None:
    table = doc.add_table(rows=1, cols=1)
    # A single-cell callout is represented by a Word table for reliable
    # background/border rendering. Mark its only row as the table header so
    # assistive technology does not treat it as an unlabeled data table.
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_borders(cell, accent, 8)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    paragraph = cell.paragraphs[0]
    set_paragraph(paragraph, after=0, line=1.15)
    set_run_font(paragraph.add_run(f"{label}  "), size=10, color=accent, bold=True)
    set_run_font(paragraph.add_run(text), size=10, color=INK)
    set_table_geometry(table, [CONTENT_DXA])
    add_text(doc, "", after=2, size=2)


def setup_document(doc: Document, running_label: str) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_ASCII
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_ASCII)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_ASCII)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_CJK)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    styles = {
        "Title": (25, NAVY, 0, 8),
        "Subtitle": (13, MUTED, 0, 16),
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (11.5, NAVY, 8, 4),
    }
    for name, (size, color, before, after) in styles.items():
        style = doc.styles[name]
        style.font.name = BODY_ASCII
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_ASCII)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_ASCII)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_CJK)
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = name.startswith("Heading") or name == "Title"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    set_paragraph(header, after=0, line=1.0)
    set_run_font(header.add_run(running_label), size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(footer.add_run("软件著作权申请文档  |  V1.0  |  "), size=8.5, color=MUTED)
    add_page_number(footer)


def create_flow_diagram(path: Path, title: str, steps: tuple[str, ...], *, accent: str) -> None:
    width, height = 1600, 780
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font_path = "/System/Library/Fonts/Hiragino Sans GB.ttc"
    title_font = ImageFont.truetype(font_path, 42)
    step_font = ImageFont.truetype(font_path, 28)
    small_font = ImageFont.truetype(font_path, 23)
    accent_rgb = tuple(int(accent[index:index + 2], 16) for index in (0, 2, 4))
    draw.text((70, 54), title, fill=(31, 77, 120), font=title_font)
    draw.text((70, 115), "系统主处理链路（管理员确认点以红色标识）", fill=(95, 99, 104), font=small_font)
    columns = 4
    rows = math.ceil(len(steps) / columns)
    box_w, box_h = 320, 150
    gap_x, gap_y = 42, 90
    start_x, start_y = 70, 210
    positions: list[tuple[int, int]] = []
    for index, step in enumerate(steps):
        row = index // columns
        col = index % columns if row % 2 == 0 else columns - 1 - index % columns
        x = start_x + col * (box_w + gap_x)
        y = start_y + row * (box_h + gap_y)
        positions.append((x, y))
        is_review = "审核" in step or "确认" in step
        outline = (166, 27, 27) if is_review else accent_rgb
        fill = (252, 232, 230) if is_review else (232, 238, 245)
        draw.rounded_rectangle((x, y, x + box_w, y + box_h), radius=18, fill=fill, outline=outline, width=4)
        number = str(index + 1)
        draw.ellipse((x + 18, y + 18, x + 62, y + 62), fill=outline)
        number_box = draw.textbbox((0, 0), number, font=small_font)
        draw.text(
            (x + 40 - (number_box[2] - number_box[0]) / 2, y + 39 - (number_box[3] - number_box[1]) / 2),
            number,
            fill="white",
            font=small_font,
            anchor="mm",
        )
        lines = wrap_cjk(step, 10)
        for line_index, line in enumerate(lines[:3]):
            bbox = draw.textbbox((0, 0), line, font=step_font)
            draw.text(
                (x + box_w / 2 - (bbox[2] - bbox[0]) / 2, y + 76 + line_index * 38),
                line,
                fill=(32, 33, 36),
                font=step_font,
            )
    for index in range(len(positions) - 1):
        x1, y1 = positions[index]
        x2, y2 = positions[index + 1]
        if abs(y2 - y1) < 10:
            if x2 > x1:
                start = (x1 + box_w + 8, y1 + box_h / 2)
                end = (x2 - 10, y2 + box_h / 2)
            else:
                start = (x1 - 8, y1 + box_h / 2)
                end = (x2 + box_w + 10, y2 + box_h / 2)
        else:
            start = (x1 + box_w / 2, y1 + box_h + 8)
            end = (x2 + box_w / 2, y2 - 10)
        draw.line((start, end), fill=(95, 99, 104), width=5)
        angle = math.atan2(end[1] - start[1], end[0] - start[0])
        arrow = []
        for offset in (2.55, -2.55):
            arrow.append((end[0] + 18 * math.cos(angle + offset), end[1] + 18 * math.sin(angle + offset)))
        draw.polygon([end, *arrow], fill=(95, 99, 104))
    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path, dpi=(180, 180))


def wrap_cjk(text: str, width: int) -> list[str]:
    return [text[index:index + width] for index in range(0, len(text), width)]


def add_cover(doc: Document, spec: SoftwareSpec) -> None:
    add_text(
        doc,
        "计算机软件著作权登记申请材料",
        size=10,
        color=RED,
        bold=True,
        before=30,
        after=26,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(spec.software_name)
    add_text(
        doc,
        spec.subtitle,
        size=14,
        color=NAVY,
        bold=True,
        after=10,
        line=1.15,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_text(
        doc,
        "软件设计说明书（含关键源程序附录）",
        size=12,
        color=MUTED,
        after=28,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_callout(
        doc,
        "软件边界",
        spec.scope_statement,
        fill=PALE_BLUE,
        accent=NAVY,
    )
    add_table(
        doc,
        ["申报项目", "拟填内容"],
        [
            ("软件简称", spec.short_name),
            ("版本号", "V1.0（建议登记版本，请按实际发布事实确认）"),
            ("著作权人", "【待申请人填写，与营业执照或身份证件一致】"),
            ("开发完成日期", "【待申请人按真实完成日期填写】"),
            ("开发方式", "独立开发【请按真实权属关系确认】"),
            ("文档形成日期", "2026年8月12日"),
        ],
        [2100, 7260],
        header_fill=NAVY,
        first_col_bold=True,
        font_size=9.5,
    )
    add_text(
        doc,
        "本说明书依据当前项目源代码编制。方括号中的申报主体、权属和日期信息必须由申请人核实后替换。",
        size=9.2,
        color=RED,
        italic=True,
        before=8,
        after=0,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    doc.add_page_break()


def add_document_control(doc: Document, spec: SoftwareSpec) -> None:
    add_heading(doc, "文档控制与申报说明", 1)
    add_table(
        doc,
        ["项目", "说明"],
        [
            ("文档用途", "用于软件著作权登记的软件说明文档，并作为申请表功能描述和源程序鉴别材料整理依据。"),
            ("取证范围", "以当前工作区的后端服务、数据模型、API、RAG 模块、前端管理页面及自动化测试为准。"),
            ("代码真实性", "附录代码由仓库源文件按原始行号自动提取；未对业务逻辑进行改写。"),
            ("权属限制", "本文不判断代码权属、合作开发关系、职务开发关系或第三方授权情况，申请人须自行核验。"),
            ("提交边界", "本文件可作为说明书主体使用，但不能替代登记申请表、身份证明、权属证明及按受理要求整理的完整源程序鉴别材料。"),
        ],
        [2100, 7260],
        header_fill=RED,
        first_col_bold=True,
    )
    add_callout(
        doc,
        "申报前必须处理",
        "替换所有【待填写】内容；确认软件名称在申请表、说明书和源程序页眉中完全一致；核对版本号、完成日期、发表状态、开发方式及著作权人信息。",
        fill=PALE_RED,
        accent=RED,
    )
    add_heading(doc, "目录", 2)
    for item in (
        "1. 软件概述与登记范围",
        "2. 建设背景、目标与用户角色",
        "3. 总体架构与业务流程",
        "4. 功能模块详细说明",
        "5. 数据、接口与状态设计",
        "6. 核心技术与创新特点",
        "7. 操作说明与典型流程",
        "8. 安全、异常与审计机制",
        "9. 测试与验收说明",
        "10. 运行环境、部署与维护",
        "附录 A. 关键源程序清单与代码摘录",
        "附录 B. 正式提交核对清单",
    ):
        # Keep the complete 12-item contents list on the document-control page.
        # A compact 1 pt rhythm avoids creating a nearly blank spillover page in
        # LibreOffice while remaining easy to scan in Word.
        add_text(doc, item, size=10.2, after=1)
def add_overview(doc: Document, spec: SoftwareSpec, bullet_id: int) -> None:
    heading = add_heading(doc, "1. 软件概述与登记范围", 1)
    heading.paragraph_format.page_break_before = True
    add_rich_text(
        doc,
        [
            ("软件全称：", {"bold": True, "color": NAVY}),
            (spec.software_name, {}),
        ],
        after=6,
        line=1.10,
    )
    add_rich_text(
        doc,
        [
            ("软件简称：", {"bold": True, "color": NAVY}),
            (spec.short_name, {}),
        ],
        after=6,
        line=1.10,
    )
    add_rich_text(
        doc,
        [
            ("建议版本：", {"bold": True, "color": NAVY}),
            ("V1.0（申请人须与申请表及源程序材料保持一致）", {}),
        ],
        after=10,
        line=1.10,
    )
    add_text(doc, spec.sections[0][1][0], after=8)
    add_callout(doc, "登记范围", spec.scope_statement, fill=PALE_GOLD, accent=GOLD)
    add_heading(doc, "1.1 软件形态", 2)
    for item in spec.sections[0][1][1:]:
        add_list_item(doc, item, bullet_id)
    add_heading(doc, "1.2 与原项目的关系", 2)
    add_text(
        doc,
        "本软件从“高校思政课 AI 智能教学辅助平台”中按独立业务目标、独立数据对象和独立处理流程划分形成。它复用平台的身份认证、课程章节和基础设施，但具备清晰的软件输入、处理状态、输出结果、管理界面和关键源程序，可作为一个功能完整的软件模块描述。",
    )


def add_background(doc: Document, spec: SoftwareSpec, bullet_id: int) -> None:
    add_heading(doc, "2. 建设背景、目标与用户角色", 1)
    paragraphs = spec.sections[1][1]
    add_text(doc, paragraphs[0])
    add_heading(doc, "2.1 建设目标", 2)
    for item in paragraphs[1:6]:
        add_list_item(doc, item, bullet_id)
    add_heading(doc, "2.2 用户角色", 2)
    add_table(
        doc,
        ["角色", "主要职责", "权限边界"],
        [
            ("管理员", "配置来源、发起任务、核验材料、确认范围、发布或归档资料", "可执行高风险审核与知识库治理操作"),
            ("教师", "查看已发布资料和变化提醒，在备课与教学中使用检索结果", "不能确认中央材料或越权访问其他教学班资料"),
            ("学生", "在课程学习、问答和复习中使用已发布资料的引用结果", "只读取与本人课程或教学班相关的正式资料"),
            ("系统任务", "执行抓取、解析、向量化、重建、重试和状态恢复", "不替代管理员作出正式发布决定"),
        ],
        [1300, 4200, 3860],
        header_fill=NAVY,
        first_col_bold=True,
    )
    add_heading(doc, "2.3 主要输入与输出", 2)
    add_text(doc, paragraphs[6])


def add_architecture(doc: Document, spec: SoftwareSpec) -> None:
    add_heading(doc, "3. 总体架构与业务流程", 1)
    diagram = ASSET_DIR / spec.diagram_name
    picture = doc.add_picture(str(diagram), width=Inches(6.5))
    picture._inline.docPr.set("descr", f"{spec.software_name}主处理链路图")
    picture._inline.docPr.set("title", "软件主处理链路")
    caption = doc.paragraphs[-1]
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(caption, after=5, line=1.0)
    add_text(doc, "图 1  软件主处理链路", size=9, color=MUTED, after=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_heading(doc, "3.1 分层架构", 2)
    add_table(
        doc,
        ["层次", "主要组件", "职责"],
        [
            ("表示层", "Vue 3、TypeScript、Element Plus", "提供任务、审核、资料、检索和状态反馈界面"),
            ("接口层", "FastAPI、Pydantic、依赖注入", "接收请求、校验参数、执行角色授权并返回统一响应"),
            ("业务层", "发现服务、资料中心服务、知识服务、AI 服务", "实现业务规则、状态流转、审核与服务编排"),
            ("检索层", "文本切分、Embedding、Chroma、分层检索", "建立向量索引、执行相似度召回与证据组合"),
            ("数据层", "SQLAlchemy、SQLite/MySQL、文件目录", "保存业务对象、正文快照、页文、分块、索引标识和审计信息"),
        ],
        [1300, 3100, 4960],
        header_fill=BLUE,
        first_col_bold=True,
    )
    add_heading(doc, "3.2 软件边界与交接", 2)
    add_text(doc, spec.sections[2][1][0])
    add_callout(doc, "关键控制点", spec.sections[2][1][1], fill=PALE_RED, accent=RED)


def add_function_details(doc: Document, spec: SoftwareSpec, bullet_id: int) -> None:
    add_heading(doc, "4. 功能模块详细说明", 1)
    content = spec.sections[3][1]
    cursor = 0
    module_number = 1
    while cursor < len(content):
        title = content[cursor]
        cursor += 1
        add_heading(doc, f"4.{module_number} {title}", 2)
        module_number += 1
        while cursor < len(content) and not content[cursor].startswith("模块："):
            add_list_item(doc, content[cursor], bullet_id)
            cursor += 1
        if cursor < len(content) and content[cursor].startswith("模块："):
            content = content[:cursor] + (content[cursor][3:],) + content[cursor + 1:]


def add_data_api_design(doc: Document, spec: SoftwareSpec) -> None:
    add_heading(doc, "5. 数据、接口与状态设计", 1)
    if "采集审核" in spec.short_name:
        rows = [
            ("AuthoritySourceRegistry", "权威来源白名单、适配器、抓取间隔、告警开关与健康状态"),
            ("DiscoveryJob", "一次后台发现任务的条件、状态、阶段、计数、异常和起止时间"),
            ("MaterialCandidate", "候选标题、来源、规范 URL、哈希、评分、建议范围和审核状态"),
            ("MaterialSnapshot", "每次成功抓取的独立正文快照、内容哈希、ETag 和解析器版本"),
            ("PolicyChange", "新旧原文证据、相似度、证据置信度、重要级别、审核与同步状态"),
        ]
        api_rows = [
            ("POST /knowledge/discovery/jobs", "创建发现任务并转入后台队列"),
            ("GET /knowledge/discovery/candidates", "按状态、来源等级读取候选材料"),
            ("POST /candidates/{id}/analyze", "执行教材关联与原文差异分析"),
            ("POST /candidates/{id}/review", "发布、驳回或标记重复"),
            ("POST /changes/{id}/review", "确认、忽略或观察政策变化证据"),
        ]
        state_text = "queued → running → completed/failed/cancelled；候选材料经历 discovered → fetched/analyzed → pending_review → published/rejected/duplicate/observed。"
    else:
        rows = [
            ("KnowledgeDocument", "文档元数据、来源类型、作用域、版本、哈希、索引集合和发布状态"),
            ("DocumentPage", "PDF 物理页或文本页、页面正文、版面尺寸和印刷页校准信息"),
            ("KnowledgeChunk", "分块内容、向量 ID、章节、页码、段落定位、锚点和索引版本"),
            ("Document*Scope", "教材、专题、教学班和知识标签的可核验适用范围"),
            ("active_index.*.json", "活动集合名、Embedding 提供方、模型、维数、指纹和清单版本"),
        ]
        api_rows = [
            ("POST /knowledge/materials", "上传中央或地方材料并建立初始索引"),
            ("PUT /materials/{id}/scopes", "确认教材、专题、教学班与知识标签范围"),
            ("POST /materials/{id}/publish", "发布材料，使其进入正式检索集合"),
            ("POST /documents/{id}/reindex", "按当前 Embedding 配置重建文档索引"),
            ("POST /knowledge/search", "按教材、专题和 Top-K 条件检索知识片段"),
        ]
        state_text = "processing → ready/failed；中央材料 pending → published → archived；已归档材料保留审计记录，但不再进入新回答。"
    add_heading(doc, "5.1 核心数据对象", 2)
    add_table(doc, ["对象", "用途"], rows, [2500, 6860], header_fill=NAVY, first_col_bold=True)
    add_heading(doc, "5.2 主要接口", 2)
    add_table(doc, ["接口", "功能"], api_rows, [3900, 5460], header_fill=BLUE, first_col_bold=True)
    add_heading(doc, "5.3 状态模型", 2)
    add_text(doc, state_text)


def add_technical_features(doc: Document, spec: SoftwareSpec, number_id: int) -> None:
    add_heading(doc, "6. 核心技术与创新特点", 1)
    for item in spec.sections[4][1]:
        add_list_item(doc, item, number_id, bold_lead=item.split("：", 1)[0] + "：" if "：" in item else None)
    add_callout(
        doc,
        "技术真实性说明",
        "上述特点均可在附录所列源文件中定位。大模型、Cross-Encoder 或 NLI 属于可选增强能力；不可用时系统保留可复现的确定性降级路径。",
        fill=PALE_GREEN,
        accent=GREEN,
    )


def add_operation_guide(doc: Document, spec: SoftwareSpec, number_id: int) -> None:
    add_heading(doc, "7. 操作说明与典型流程", 1)
    add_heading(doc, "7.1 管理员典型操作", 2)
    for item in spec.sections[5][1]:
        add_list_item(doc, item, number_id)
    add_heading(doc, "7.2 页面反馈与结果判断", 2)
    add_table(
        doc,
        ["界面状态", "用户可见反馈", "允许操作"],
        [
            ("等待/处理中", "显示当前阶段、处理计数、开始时间和后台运行提示", "查看详情、离开页面、必要时取消"),
            ("待审核/待发布", "展示原文、范围、证据、评分或索引就绪状态", "确认、修正范围、驳回、观察"),
            ("成功", "显示已发布、已同步或检索命中结果及来源定位", "继续管理、归档、再次检索"),
            ("失败", "显示人类可读错误、失败阶段和可重试状态", "重试、修正配置、检查原文"),
        ],
        [1600, 4800, 2960],
        header_fill=NAVY,
        first_col_bold=True,
    )


def add_security(doc: Document, spec: SoftwareSpec, bullet_id: int) -> None:
    add_heading(doc, "8. 安全、异常与审计机制", 1)
    for item in spec.sections[6][1]:
        add_list_item(doc, item, bullet_id)
    add_callout(
        doc,
        "人工责任边界",
        "系统输出的相关性、差异、范围和提醒建议均为辅助信息；涉及中央材料发布、政策变化确认、教材范围调整或教师提醒时，必须由具备权限的管理员核对原文后决定。",
        fill=PALE_RED,
        accent=RED,
    )


def add_tests(doc: Document, spec: SoftwareSpec) -> None:
    add_heading(doc, "9. 测试与验收说明", 1)
    if "采集审核" in spec.short_name:
        rows = [
            ("来源安全", "拒绝非 HTTPS、跨白名单域名、内网或本机地址", "已由来源适配器与发现服务测试覆盖"),
            ("任务可靠性", "任务失败可见、重试可用、原子认领避免重复执行", "已由发现任务状态测试覆盖"),
            ("候选治理", "正文哈希/规范 URL 去重，审核状态合法流转", "已由 authority_discovery 测试覆盖"),
            ("教材关联", "BM25、向量、词项和 RRF 召回可降级运行", "已由匹配服务与评估脚本覆盖"),
            ("变化证据", "章节级来源、证据置信度和人工确认分离", "已由政策变化与通知测试覆盖"),
        ]
    else:
        rows = [
            ("文档入库", "支持 PDF/TXT/Markdown，拒绝伪造扩展名和空正文", "已由 knowledge 测试覆盖"),
            ("索引建立", "分块、向量 ID、页码定位和数据库记录一致", "已由 knowledge 与 citation 测试覆盖"),
            ("检索回答", "检索结果可供 AI 使用并返回来源定位", "已由 knowledge/ai 测试覆盖"),
            ("模型切换", "维数或指纹不一致时禁止复用旧集合", "已由 embedding_index 测试覆盖"),
            ("重建与删除", "重建可恢复；已发布资料须先归档再处理", "已由 knowledge/material_center 测试覆盖"),
        ]
    add_table(
        doc,
        ["测试域", "主要验收点", "项目证据"],
        rows,
        [1600, 4800, 2960],
        header_fill=BLUE,
        first_col_bold=True,
    )
    add_heading(doc, "9.1 建议登记前验收", 2)
    add_text(
        doc,
        "在正式提交前，建议使用申请版本执行后端自动化测试、前端类型检查与生产构建，并保留通过日期、测试人员、环境和失败修复记录。若申请版本与当前工作区不同，应重新生成代码附录。",
    )


def add_runtime(doc: Document, spec: SoftwareSpec) -> None:
    add_heading(doc, "10. 运行环境、部署与维护", 1)
    add_table(
        doc,
        ["类别", "推荐配置"],
        [
            ("服务端", "Python 3.11+，FastAPI，SQLAlchemy；开发可用 SQLite，部署可用 MySQL"),
            ("前端", "Node.js 24+，Vue 3，TypeScript，Vite，Element Plus"),
            ("向量检索", "Chroma；开发可用 256 维确定性模拟 Embedding，生产接入兼容 Embedding 服务"),
            ("操作系统", "Linux/macOS/Windows 开发环境；生产建议 Linux 容器或受控服务器"),
            ("浏览器", "支持现代 Chromium、Firefox、Safari 的近期稳定版本"),
            ("存储", "数据库、原始资料目录、Chroma 持久化目录应分别备份并控制访问权限"),
        ],
        [1800, 7560],
        header_fill=NAVY,
        first_col_bold=True,
    )
    add_heading(doc, "10.1 部署与配置", 2)
    add_text(doc, spec.sections[7][1][0])
    add_heading(doc, "10.2 维护与备份", 2)
    add_text(doc, spec.sections[7][1][1])
    add_heading(doc, "10.3 已知边界", 2)
    add_text(doc, spec.sections[7][1][2])


def file_digest(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def add_code_block(doc: Document, source_path: Path, start: int, end: int) -> None:
    lines = source_path.read_text(encoding="utf-8").splitlines()
    if start < 1 or end > len(lines) or start > end:
        raise ValueError(f"Invalid source range: {source_path}:{start}-{end}")
    for line_number in range(start, end + 1):
        paragraph = doc.add_paragraph()
        set_paragraph(paragraph, before=0, after=0, line=1.0)
        paragraph.paragraph_format.left_indent = Inches(0.08)
        paragraph.paragraph_format.right_indent = Inches(0.04)
        paragraph.paragraph_format.widow_control = False
        shade_paragraph(paragraph, PALE_GRAY)
        prefix = f"{line_number:04d}  "
        set_run_font(
            paragraph.add_run(prefix),
            size=7.2,
            color=MUTED,
            ascii_font=CODE_ASCII,
            cjk_font=CODE_CJK,
        )
        set_run_font(
            paragraph.add_run(lines[line_number - 1] or " "),
            size=7.2,
            color=INK,
            ascii_font=CODE_ASCII,
            cjk_font=CODE_CJK,
        )
    add_text(doc, "", after=3, size=2)


def add_code_appendix(doc: Document, spec: SoftwareSpec) -> None:
    heading = add_heading(doc, "附录 A. 关键源程序清单与代码摘录", 1)
    heading.paragraph_format.page_break_before = True
    add_text(
        doc,
        "本附录按当前仓库原始行号自动提取关键源程序，用于说明核心实现和辅助整理源程序鉴别材料。代码摘录不是完整源程序提交件；正式登记时应按受理要求另行整理申请版本的完整源程序页。",
    )
    inventory = []
    for index, excerpt in enumerate(spec.code_excerpts, start=1):
        source_path = ROOT / excerpt.source
        ranges = "、".join(f"{start}-{end}" for start, end in excerpt.ranges)
        inventory.append((index, excerpt.source, ranges, excerpt.title))
    add_table(
        doc,
        ["序号", "源文件", "原始行号", "核心功能"],
        inventory,
        [700, 3900, 1500, 3260],
        header_fill=RED,
        font_size=8.2,
    )
    for index, excerpt in enumerate(spec.code_excerpts, start=1):
        source_path = ROOT / excerpt.source
        add_heading(doc, f"A.{index} {excerpt.title}", 2)
        add_text(doc, excerpt.purpose, size=9.6, after=4)
        add_text(
            doc,
            f"源文件：{excerpt.source}  |  SHA-256：{file_digest(source_path)[:16]}…",
            size=8.5,
            color=MUTED,
            after=4,
        )
        for part_index, (start, end) in enumerate(excerpt.ranges, start=1):
            if len(excerpt.ranges) > 1:
                add_text(
                    doc,
                    f"代码片段 {part_index}/{len(excerpt.ranges)}：原始行 {start}-{end}",
                    size=8.5,
                    color=NAVY,
                    bold=True,
                    after=3,
                    keep_with_next=True,
                )
            add_code_block(doc, source_path, start, end)


def add_submission_checklist(doc: Document, spec: SoftwareSpec, bullet_id: int) -> None:
    add_heading(doc, "附录 B. 正式提交核对清单", 1)
    for item in (
        f"申请表、说明书、源程序页眉中的软件名称统一为“{spec.software_name}”，版本号统一为 V1.0 或申请人确认的实际版本。",
        "删除或替换全部【待填写】内容，著作权人名称与证件、营业执照或事业单位法人证书完全一致。",
        "核实开发完成日期、首次发表日期、发表状态、开发方式和权利取得方式，不根据本文示例直接填写。",
        "按受理要求另行准备源程序鉴别材料；源程序应来自与本说明书相同的申请版本，并连续编排页码和行号。",
        "检查第三方依赖、开源许可、合作开发、委托开发、职务开发和素材授权，必要时准备相应权属证明。",
        "保存申请版本的代码归档、数据库结构、部署文件、测试报告和形成时间证据，以备补正或权属核验。",
        "在 Word 中最终更新目录/页码并进行人工通读，确认无内部账号、密钥、真实个人隐私、测试密码或无关项目材料。",
    ):
        add_list_item(doc, item, bullet_id)
    add_callout(
        doc,
        "交付结论",
        "完成上述核对并补齐主体信息后，本说明书可直接作为软件著作权申请的软件文档主体；申请表、身份证明/主体资格证明、权属证明及源程序鉴别材料仍须按登记机构当期要求一并提交。",
        fill=PALE_GOLD,
        accent=GOLD,
    )


def build_document(spec: SoftwareSpec) -> Path:
    doc = Document()
    setup_document(doc, spec.running_label)
    doc.core_properties.title = spec.software_name
    doc.core_properties.subject = "计算机软件著作权登记申请材料"
    doc.core_properties.author = "待申请人填写"
    doc.core_properties.comments = "依据当前项目源代码自动编制，申报主体信息须人工核对。"
    bullet_id = add_numbering(doc, bullet=True)
    number_id = add_numbering(doc, bullet=False)

    add_cover(doc, spec)
    add_document_control(doc, spec)
    add_overview(doc, spec, bullet_id)
    add_background(doc, spec, bullet_id)
    add_architecture(doc, spec)
    add_function_details(doc, spec, bullet_id)
    add_data_api_design(doc, spec)
    add_technical_features(doc, spec, number_id)
    add_operation_guide(doc, spec, number_id)
    add_security(doc, spec, bullet_id)
    add_tests(doc, spec)
    add_runtime(doc, spec)
    add_code_appendix(doc, spec)
    add_submission_checklist(doc, spec, bullet_id)

    output = OUTPUT_DIR / spec.output_name
    doc.save(output)
    return output


DISCOVERY_SPEC = SoftwareSpec(
    output_name="软著申请材料-权威新闻智能采集审核与知识库对比系统-V1.0.docx",
    software_name="高校思政课权威新闻智能采集审核与知识库对比系统",
    short_name="权威新闻采集审核与知识库对比系统",
    subtitle="面向最新权威新闻与政策材料的自动发现、审核、去重和既有知识对比",
    scope_statement=(
        "本软件从权威来源配置开始，完成网页线索发现、正文抓取、快照留存、重复识别、教材专题关联、"
        "新旧原文差异证据生成及管理员审核。其正式输出是“经人工确认的候选材料、适用范围和变化证据”；"
        "未经确认的数据不进入正式 RAG 知识库。"
    ),
    running_label="权威新闻智能采集审核与知识库对比系统  |  软件设计说明书",
    diagram_name="discovery_flow.png",
    diagram_steps=("配置权威来源", "创建发现任务", "抓取并提取正文", "规范化与去重", "关联教材专题", "生成差异证据", "管理员审核确认", "输出已审核材料"),
    sections=(
        (
            "overview",
            (
                "本软件面向高校思政课资料治理场景，自动巡检中国政府网、教育部门和权威媒体等白名单来源，识别最新新闻、政策文件和重要讲话，形成可核验正文快照，并与教材章节及既有中央材料进行关联和差异比较。",
                "采用浏览器/服务器架构，管理员通过网页完成来源、任务、候选和差异审核。",
                "后端以独立后台任务执行抓取和分析，页面关闭后任务继续运行。",
                "支持 HTML 列表、RSS、Sitemap 和单篇原文四类来源入口。",
                "候选材料、正文快照、差异证据与审核人信息分别保存，具备追溯能力。",
                "软件不以模型结论直接改变知识库，正式发布必须经过管理员确认。",
            ),
        ),
        (
            "background",
            (
                "高校思政课需要及时吸收权威新闻与政策表述，但人工逐站检索、复制、去重和教材定位耗时，且容易遗漏原文来源、发布时间与历史表述。软件以白名单、原文快照和人工审核为基础，把“发现线索”转化为“可核验的更新证据”。",
                "持续发现白名单权威来源中的最新材料，并记录来源健康状态。",
                "安全提取正文与结构化元数据，降低导航、推荐链接和页面噪声干扰。",
                "通过规范 URL、内容哈希和议题聚类减少重复审核。",
                "使用教材章节多路召回定位可能受影响的课程和专题。",
                "生成新旧原文句段级证据，并把算法建议与人工结论分开保存。",
                "输入包括来源配置、关键词、日期范围和已有教材/知识文档；输出包括发现任务、候选材料、正文快照、关联建议、差异证据、审核结果和同步状态。",
            ),
        ),
        (
            "boundary",
            (
                "本软件读取课程章节和已发布知识文档作为比较基线，但不负责最终向量索引的构建和 RAG 问答。审核通过后，通过资料中心服务把正文、来源元数据、教材范围和变化状态交给“知识库智能合并与 RAG 增量更新系统”。",
                "发现、相关性、重要度、教材关联和差异类型均为辅助判断；只有管理员的发布或变化确认操作才构成正式知识更新指令。",
            ),
        ),
        (
            "functions",
            (
                "来源白名单与适配器管理",
                "校验来源必须属于配置域名，默认要求公开 HTTPS，并阻止访问本机或内网地址。",
                "为中国政府网、教育部、求是网等来源配置专用正文标记、噪声标记和详情页路径规则。",
                "保存抓取间隔、请求间隔、全文权限、提醒开关、连续失败次数和最近错误。",
                "模块：后台发现任务",
                "按关键词、日期范围和来源创建任务，限制排队数量和并发任务数。",
                "通过数据库条件更新原子认领任务，避免多进程重复执行同一任务。",
                "服务重启后将遗留 running 任务恢复为 queued，管理员可查看进度、失败原因并重试。",
                "模块：网页抓取与正文抽取",
                "读取 HTML 列表、RSS、Sitemap 或单篇原文，优先处理标题命中线索。",
                "限制重定向次数、响应体大小和请求重试次数，记录最终 URL、ETag 和 Last-Modified。",
                "使用来源适配器识别 article/main/content 区域并过滤脚本、导航、页脚和推荐模块。",
                "模块：快照、质量评分与去重",
                "为每次成功抓取保存独立正文快照和 SHA-256 内容哈希。",
                "规范化 URL，移除常见跟踪参数并统一主机名、路径和查询参数。",
                "结合 URL 哈希、正文哈希、标题和议题相似性识别重复内容或旁证材料。",
                "模块：教材关联与重要度分析",
                "从标题和政策段落构造压缩查询，避免长篇通用表述淹没主题。",
                "使用 BM25、向量召回和中文二元词项相似度形成多路排名，再使用 RRF 融合。",
                "可选 BGE Cross-Encoder 和受约束大模型复核；不可用时退回确定性门槛。",
                "模块：知识库对比与差异证据",
                "仅在建议章节和已发布中央材料范围内选择旧材料，防止跨专题误比较。",
                "从新旧正文选取政策句段，使用主题相关性、字面相似度和可选 NLI 识别新增、调整、强化或解释更新。",
                "分别保存章节关联置信度和句段证据置信度；只有双重阈值满足时建议提醒。",
                "模块：人工审核与结果输出",
                "管理员可发布、驳回、标记重复、观察、批量处理或删除未发布候选。",
                "发布前必须补全发布机构、发布日期，并至少确认一本教材或一个专题。",
                "政策变化可确认、忽略或观察；确认后由后续知识更新系统执行索引同步。",
            ),
        ),
        (
            "features",
            (
                "安全白名单抓取：将公开 HTTPS、域名边界、地址解析、重定向和体积限制统一到抓取入口。",
                "来源适配正文抽取：按站点结构提取内容和元数据，并以通用解析器作为兼容回退。",
                "双哈希去重：规范 URL 的固定长度哈希解决数据库长索引限制，正文哈希识别同内容多地址。",
                "多路教材召回：BM25、向量和中文词项通道经 RRF 融合，兼顾可复现性与语义召回。",
                "双置信度门控：章节关联和新旧证据分别评分，避免把“相关文章”直接误判为“政策变化”。",
                "快照与证据可追溯：正文快照、解析器版本、新旧摘录、来源 URL 和审核记录相互关联。",
                "人机权限分离：AI 只给建议，管理员确认后才输出正式材料或变化指令。",
                "可靠后台调度：并发上限、队列容量、原子认领、失败重试和重启恢复共同保证任务可管理。",
            ),
        ),
        (
            "operations",
            (
                "进入“资料动态”，检查或新增权威来源，确认域名、入口、来源等级和抓取间隔。",
                "设置关键词、日期范围和来源，提交发现任务；任务进入后台后可离开页面。",
                "在任务列表查看已发现、已抓取、已去重、待审核、过滤和失败数量。",
                "打开候选材料，核对标题、发布机构、发布日期、原文链接、正文快照和质量评分。",
                "执行“全教材关联与原文差异分析”，查看建议课程、专题和新旧句段证据。",
                "按原文确认候选的教材范围、知识标签和材料状态；可发布、驳回、标记重复或加入观察。",
                "对政策变化执行确认、忽略或观察；已确认结果进入知识库更新交接流程。",
            ),
        ),
        (
            "security",
            (
                "来源安全：公开 HTTPS、白名单域名、内网地址阻断和跨域降级限制降低 SSRF 风险。",
                "访问控制：所有发现、候选和变化接口均要求管理员角色。",
                "输入约束：关键词、来源数量、批量数量、文本长度、日期范围和 URL 长度均有限制。",
                "资源控制：请求体大小、每日抓取量、单来源链接数、并发数和队列数受配置约束。",
                "异常可见：来源错误、任务错误、正文抽取失败和分析失败分别计数并记录。",
                "审计追踪：候选和变化保存审核人、审核时间、审核备注、快照和来源 URL。",
                "数据保护：已发布候选不能直接从候选池删除，必须在资料中心执行归档。",
            ),
        ),
        (
            "runtime",
            (
                "部署时配置数据库、允许的来源、抓取并发、调度开关和每日限额。默认来源可初始化，但生产环境应由管理员核验站点条款、公开访问条件和抓取频率。",
                "定期备份数据库和正文快照；升级前执行数据库迁移。对连续失败来源、堆积任务、低质量正文和待审核数量建立运维检查。",
                "网页结构变化可能降低正文抽取质量；部分站点限制自动访问。可选模型不可用时精排能力下降，但确定性召回、审核和追溯流程仍可运行。",
            ),
        ),
    ),
    code_excerpts=(
        CodeExcerpt(
            "白名单 URL 校验与规范化",
            "统一检查 HTTPS、公开网络地址和域名边界，并移除跟踪参数形成稳定 URL。",
            "backend/app/services/authority_discovery_service.py",
            ((113, 143),),
        ),
        CodeExcerpt(
            "详情页抓取、正文解析与线索筛选",
            "按来源适配器抓取详情页、提取有效正文，并在列表线索中执行域名过滤和规范 URL 去重。",
            "backend/app/services/authority_discovery_service.py",
            ((341, 388),),
        ),
        CodeExcerpt(
            "来源适配器的正文和元数据抽取",
            "识别正文容器、排除噪声区域，并从 Meta、H1 和可见文本提取标题、发布者与日期。",
            "backend/app/services/authority_source_adapters.py",
            ((82, 166),),
        ),
        CodeExcerpt(
            "审核通过后的材料交接",
            "校验发布条件，把最新正文快照、来源信息和确认范围交给资料中心，并保留候选审核记录。",
            "backend/app/services/authority_discovery_service.py",
            ((711, 764),),
        ),
        CodeExcerpt(
            "BM25、向量与词项召回的 RRF 融合",
            "组合三个检索通道，计算确定性置信度并生成管理员可核对的专题关联建议。",
            "backend/app/services/authority_discovery_service.py",
            ((949, 1056),),
        ),
        CodeExcerpt(
            "新旧原文差异证据与双阈值门控",
            "对候选句段执行主题匹配、相似度比较和可选 NLI，并生成带证据置信度的变化记录。",
            "backend/app/services/authority_discovery_service.py",
            ((1191, 1215), (1237, 1292)),
        ),
        CodeExcerpt(
            "任务原子认领与失败可见",
            "通过数据库条件更新认领排队任务，防止并发 Web 进程重复处理同一任务。",
            "backend/app/services/authority_discovery_service.py",
            ((1404, 1450),),
        ),
    ),
)


RAG_SPEC = SoftwareSpec(
    output_name="软著申请材料-知识库智能合并与RAG增量更新系统-V1.0.docx",
    software_name="高校思政课知识库智能合并与 RAG 增量更新系统",
    short_name="知识库智能合并与 RAG 增量更新系统",
    subtitle="审核材料的正式入库、向量索引构建、分层检索、增量重建与安全切换",
    scope_statement=(
        "本软件接收经审核的教材、中央材料和地方材料，完成原文存档、内容哈希校验、适用范围确认、"
        "页内切分、Embedding、Chroma 向量写入、关系数据同步、分层检索和 RAG 上下文构建；"
        "并在材料更新或模型切换时执行可重试的增量重建与活动索引安全切换。"
    ),
    running_label="知识库智能合并与 RAG 增量更新系统  |  软件设计说明书",
    diagram_name="rag_flow.png",
    diagram_steps=("接收审核材料", "存档并校验哈希", "确认资料作用域", "页内切分与定位", "生成语义向量", "写入 Chroma 与数据库", "分层检索与引用", "更新重建并确认"),
    sections=(
        (
            "overview",
            (
                "本软件面向高校思政课知识库构建与持续更新场景，把教材正文、中央权威材料和地方教学材料转换为可检索、可定位、可发布和可回退的 RAG 知识资产，并为 AI 问答与备课提供来源清晰的检索上下文。",
                "采用浏览器/服务器架构，管理员或授权教师通过资料中心完成导入、范围确认、发布、归档和重建。",
                "支持 PDF、TXT、Markdown 以及公开 HTTPS 网页正文存档。",
                "同时保存原始文件、页文、文本分块、向量 ID、索引版本和来源元数据。",
                "中央材料、教材、地方材料分层治理并按权限进入检索。",
                "Embedding 模型、维数和协议以指纹隔离，禁止不同维数写入同一活动集合。",
            ),
        ),
        (
            "background",
            (
                "通用大模型无法保证回答符合课程教材和最新权威原文。知识库需要持续接收审核材料，并解决文件解析、细粒度定位、向量维数兼容、旧版本归档、检索配额和引用核验等工程问题。本软件把这些步骤组合为可恢复的 RAG 构建与更新流程。",
                "将审核通过的原文转化为带教材、专题、页码和来源信息的知识分块。",
                "同时维护关系数据库与向量数据库的一致性，失败时进入明确可重试状态。",
                "按中央材料、教材正文和地方材料分别召回，避免权威加权覆盖相关性门槛。",
                "支持文档重建、版本替代、归档和 Embedding 模型安全切换。",
                "为 AI 回答返回可核验的来源标题、原文 URL、PDF 页码和段落位置。",
                "输入包括审核材料、原文件、来源元数据、教材/专题/教学班范围和 Embedding 配置；输出包括正式知识文档、页文、知识分块、向量集合、检索片段、引用元数据和同步状态。",
            ),
        ),
        (
            "boundary",
            (
                "本软件的上游是资料采集审核系统或人工资料导入，输入必须包含可存档正文和经权限校验的适用范围。下游是 AI 学习助手、备课 Agent、学生笔记和知识检索接口；下游只能读取状态为 ready、已发布且在有效作用域内的材料。",
                "发布状态与向量就绪状态同时成立才允许资料进入正式检索；模型切换时必须验证活动索引清单和向量维数，不能把新向量混写到旧集合。",
            ),
        ),
        (
            "functions",
            (
                "多源资料导入与存档",
                "支持文本型 PDF、TXT、Markdown 和公开 HTTPS 网页；验证扩展名、PDF 文件头、文件大小和正文非空。",
                "对 URL 导入执行公开地址校验和重定向控制，从网页结构化信息补充标题、发布机关和日期。",
                "计算 SHA-256 内容哈希，阻止同类型相同正文重复入库。",
                "模块：资料分类、范围和发布",
                "按中央材料、教材正文、地方材料和待分类四种类型管理。",
                "确认教材、专题、教学班和知识标签范围；教师地方材料必须限定到本人教学班。",
                "发布时校验原文件哈希、来源信息、索引就绪状态和必要作用域；旧版本可自动归档。",
                "模块：原文解析与精确切分",
                "PDF 按物理页提取文本，TXT/Markdown 支持 UTF-8 和 GB18030。",
                "在单页内按段落和配置长度切分，保留 PDF 起止页、段落序号、首尾锚点和章节信息。",
                "扫描版 PDF 无可提取文字时明确提示先执行 OCR，避免生成空索引。",
                "模块：Embedding 与向量写入",
                "支持确定性模拟 Embedding、OpenAI 兼容服务和 DashScope 等提供方。",
                "对模型返回向量逐批校验实际维数，不一致时中止写入并提示重建。",
                "以 provider、model、dimensions 和 rag-v2 协议生成指纹与独立 collection 名。",
                "模块：关系数据与 Chroma 一致性",
                "每个分块同时写入 Chroma 和 KnowledgeChunk，向量 ID 采用 document-{id}-chunk-{index} 稳定格式。",
                "向量元数据保存来源、材料层级、教材、专题、页码、段落、锚点和可读位置标签。",
                "失败时回滚数据库、清理残留向量和分块，并把文档标记为 failed。",
                "模块：分层 RAG 检索",
                "中央、教材和地方材料独立召回，先通过相似度阈值，再添加有限权威加分。",
                "单次回答最多 2 条中央材料、至少 2 条教材依据、最多 1 条地方材料，剩余位置由教材补足。",
                "检索结果进入 AI 提示词前再次核验文档仍存在、状态就绪、作用域有效和页码可用。",
                "模块：增量重建与安全切换",
                "材料更新时删除该文档旧向量和旧分块，重新解析、向量化并更新索引版本。",
                "全量重建可在新 collection 完成数量与覆盖率校验后原子切换活动清单。",
                "活动清单保留 previous 备份，并按 Embedding 指纹保存独立指针，支持安全回退。",
            ),
        ),
        (
            "features",
            (
                "双存储一致性：原文件、关系数据和 Chroma 向量用稳定文档/分块标识关联。",
                "精确引用定位：页内切分保留物理页、印刷页、段落序号和首尾文本锚点。",
                "Embedding 指纹隔离：提供方、模型、维数和索引协议共同决定集合，阻止维数污染。",
                "活动索引清单：以 JSON 清单记录活动集合和指纹，临时文件写入后通过 os.replace 原子替换。",
                "分层证据配额：不同材料层独立召回，权威等级只在已达相关性阈值的结果间调序。",
                "发布与索引双门控：资料原文、范围、哈希和向量全部就绪后才能进入正式 RAG。",
                "可恢复重建：失败事务回滚并清理残留，文档进入 failed 后可再次执行 reindex。",
                "版本化更新：新资料可指向 supersedes_document_id，发布新版本时归档旧版本而不丢失审计链。",
            ),
        ),
        (
            "operations",
            (
                "进入“资料中心”，选择中央材料、教材正文或地方材料页签。",
                "上传 PDF/TXT/Markdown 或提交公开 HTTPS 原文；填写标题、发布机关、日期、版本和访问策略。",
                "查看智能关联建议，确认教材、专题、教学班和知识标签范围。",
                "确认文档状态为 ready、分块数量正确、原文可访问后执行发布。",
                "在知识检索接口输入问题、教材、专题和 Top-K，核对返回片段、相关度、页码和来源。",
                "材料正文或 Embedding 配置变化时执行重新索引；若切换模型，先在新集合重建并验证后再激活。",
                "不再参与回答的资料先执行归档；删除前确认它不是当前教材版本或活动发布资料。",
            ),
        ),
        (
            "security",
            (
                "角色授权：教材和资料管理仅限教师/管理员，中央材料发布仅限管理员。",
                "班级隔离：地方材料按所有者和教学班作用域检查，学生不能通过猜测 ID 读取其他班资料。",
                "URL 安全：网页导入只允许公开 HTTPS，解析后固定公共 IP 并限制重定向与响应大小。",
                "文件安全：验证扩展名、PDF 魔数、大小、正文和内容哈希，不执行上传文件。",
                "索引安全：活动清单版本、Embedding 配置和集合实际维数三重校验。",
                "删除保护：当前教材和活动发布资料不能直接删除，必须先切换版本或归档。",
                "异常恢复：入库或重建失败时清理向量和分块，防止半完成数据参与检索。",
            ),
        ),
        (
            "runtime",
            (
                "部署时配置上传目录、Chroma 持久化目录、活动 collection、Embedding 提供方、模型、维数、相似度阈值和 Top-K。真实语义检索应配置生产 Embedding 服务。",
                "数据库、上传原文、Chroma 目录和活动索引清单必须作为一个恢复单元备份。模型切换前在新 collection 完成全量构建和覆盖率核验，保留上一份活动清单。",
                "扫描 PDF 当前不内置 OCR；模拟 Embedding 仅用于流程验证；单文档同步入库可能受外部 Embedding 延迟影响。生产部署应监控向量化失败、维数错误、索引规模和检索阈值。",
            ),
        ),
    ),
    code_excerpts=(
        CodeExcerpt(
            "文档解析与页内切分",
            "从 PDF/TXT/Markdown 提取正文，并在页内切分时保留物理页、段落序号与文本锚点。",
            "backend/app/rag/document_loader.py",
            ((19, 44),),
        ),
        CodeExcerpt(
            "页内分块定位元数据",
            "把每个页面的文本切分为独立分块，避免机械合并无关相邻页面。",
            "backend/app/rag/text_splitter.py",
            ((6, 28),),
        ),
        CodeExcerpt(
            "知识文档入库、向量化和失败恢复",
            "校验文件、保存原文和页文，将精确分块写入向量索引及关系数据库，并在失败时清理半成品。",
            "backend/app/services/knowledge_service.py",
            ((79, 115), (131, 221)),
        ),
        CodeExcerpt(
            "带引用定位的 Chroma 向量写入",
            "构造稳定向量 ID，并把页码、段落、章节、来源和可读位置标签写入向量元数据。",
            "backend/app/rag/vector_store.py",
            ((212, 246),),
        ),
        CodeExcerpt(
            "Embedding 配置指纹与维数校验",
            "为索引生成配置指纹，并在外部服务返回向量后逐条校验实际维数。",
            "backend/app/rag/embeddings.py",
            ((17, 30), (70, 105), (150, 175)),
        ),
        CodeExcerpt(
            "活动索引清单校验与原子切换",
            "验证 provider/model/dimensions/fingerprint，并通过临时文件和 os.replace 原子更新活动索引指针。",
            "backend/app/rag/vector_store.py",
            ((71, 103), (106, 149)),
        ),
        CodeExcerpt(
            "文档增量重建",
            "删除旧文档向量、页文和分块，按当前 Embedding 配置重新建立索引，并更新文档状态。",
            "backend/app/services/knowledge_service.py",
            ((255, 308),),
        ),
        CodeExcerpt(
            "中央、教材和地方材料分层检索",
            "各材料层独立召回并应用相关性阈值、有限权威加分和证据配额。",
            "backend/app/rag/retriever.py",
            ((14, 82),),
        ),
        CodeExcerpt(
            "发布、版本归档与确认变化同步",
            "材料发布时校验原文哈希和作用域并归档旧版本；已确认变化通过统一知识服务重建索引。",
            "backend/app/services/material_center_service.py",
            ((443, 480),),
        ),
        CodeExcerpt(
            "政策变化触发 RAG 索引更新",
            "在候选材料已发布后按当前 Embedding profile 重建文档，并保存同步成功、等待或失败状态。",
            "backend/app/services/authority_discovery_service.py",
            ((1338, 1372),),
        ),
    ),
)


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    if ASSET_DIR.exists():
        shutil.rmtree(ASSET_DIR)
    ASSET_DIR.mkdir(parents=True)
    create_flow_diagram(
        ASSET_DIR / DISCOVERY_SPEC.diagram_name,
        "权威新闻发现、审核与知识库对比流程",
        DISCOVERY_SPEC.diagram_steps,
        accent=BLUE,
    )
    create_flow_diagram(
        ASSET_DIR / RAG_SPEC.diagram_name,
        "知识库合并、RAG 构建与增量更新流程",
        RAG_SPEC.diagram_steps,
        accent=GREEN,
    )
    outputs = [build_document(DISCOVERY_SPEC), build_document(RAG_SPEC)]
    shutil.rmtree(ASSET_DIR)
    for output in outputs:
        print(output)


if __name__ == "__main__":
    main()
